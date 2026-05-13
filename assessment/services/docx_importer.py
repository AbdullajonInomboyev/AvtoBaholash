"""
AvtoBaholash — Word (.docx) test import servisi.
Qo'llab-quvvatlanadi: oddiy matn, OMML formulalar, inline rasmlar, LaTeX, aralash kontent.
"""
import os
import re
import random
import tempfile
import shutil
from html import escape

from django.core.files.base import ContentFile

try:
    from docx import Document
    from lxml import etree
    HAS_LXML = True
except ImportError:
    HAS_LXML = False


XML_NS = {
    "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "m":   "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v":   "urn:schemas-microsoft-com:vml",
}


# OMML → LaTeX oddiy konvertor (asosiy elementlar)
def _omml_to_latex(omml_element):
    """Microsoft Equation OMML XML dan LaTeX matn yaratadi"""
    if omml_element is None:
        return ""
    
    def walk(el):
        tag = etree.QName(el).localname if el.tag else ''
        children = list(el)

        # Matn run
        if tag == 't':
            return el.text or ""

        # Run: barcha child larni birlashtir
        if tag == 'r':
            return "".join(walk(c) for c in children)

        # Fraksiya (kasr)
        if tag == 'f':
            num = next((walk(c) for c in children if etree.QName(c).localname == 'num'), '')
            den = next((walk(c) for c in children if etree.QName(c).localname == 'den'), '')
            return f"\\frac{{{num}}}{{{den}}}"

        # Daraja / pastki/ustki indeks
        if tag == 'sSup':
            base = next((walk(c) for c in children if etree.QName(c).localname == 'e'), '')
            sup  = next((walk(c) for c in children if etree.QName(c).localname == 'sup'), '')
            return f"{{{base}}}^{{{sup}}}"
        if tag == 'sSub':
            base = next((walk(c) for c in children if etree.QName(c).localname == 'e'), '')
            sub  = next((walk(c) for c in children if etree.QName(c).localname == 'sub'), '')
            return f"{{{base}}}_{{{sub}}}"

        # Kvadrat ildiz
        if tag == 'rad':
            e   = next((walk(c) for c in children if etree.QName(c).localname == 'e'), '')
            deg = next((walk(c) for c in children if etree.QName(c).localname == 'deg'), '')
            if deg:
                return f"\\sqrt[{deg}]{{{e}}}"
            return f"\\sqrt{{{e}}}"

        # Integral, summa va boshqa
        if tag == 'nary':
            chr_el = next((c for c in children if etree.QName(c).localname == 'naryPr'), None)
            chr_val = '\\int'
            if chr_el is not None:
                ch_node = chr_el.find(f"{{{XML_NS['m']}}}chr")
                if ch_node is not None:
                    val = ch_node.get(f"{{{XML_NS['m']}}}val", '')
                    chr_val = {'∑': '\\sum', '∏': '\\prod', '∫': '\\int'}.get(val, '\\int')
            sub = next((walk(c) for c in children if etree.QName(c).localname == 'sub'), '')
            sup = next((walk(c) for c in children if etree.QName(c).localname == 'sup'), '')
            e   = next((walk(c) for c in children if etree.QName(c).localname == 'e'), '')
            res = chr_val
            if sub: res += f"_{{{sub}}}"
            if sup: res += f"^{{{sup}}}"
            res += f" {e}"
            return res

        # Default: barcha child larni birlashtirish
        text = "".join(walk(c) for c in children)
        if not text and el.text:
            return el.text
        return text

    try:
        return walk(omml_element)
    except Exception:
        # Xato bo'lsa — oddiy matn
        try:
            return " ".join(t.strip() for t in omml_element.itertext() if t.strip())
        except Exception:
            return ""


class AvtoBaholashDocxParser:
    def __init__(self, file_path):
        self.doc           = Document(file_path)
        self.related_parts = self.doc.part.related_parts
        self.errors        = []

    def _extract_paragraph(self, para):
        """Paragraphdan blocks chiqaradi"""
        blocks = []
        for child in para._element.iterchildren():
            tag = etree.QName(child).localname

            if tag == 'r':
                # Matn
                texts = child.findall(".//w:t", namespaces=XML_NS)
                t = "".join(x.text or "" for x in texts).strip()
                if t:
                    blocks.append({"type": "text", "value": t})
                # Inline OMML
                for omath in child.findall(".//m:oMath", namespaces=XML_NS):
                    latex = _omml_to_latex(omath)
                    if latex:
                        blocks.append({"type": "formula", "latex": latex})
                # Rasmlar
                for blip in child.findall(".//a:blip", namespaces=XML_NS):
                    rid = blip.get(f'{{{XML_NS["r"]}}}embed')
                    if rid and rid in self.related_parts:
                        part = self.related_parts[rid]
                        blocks.append({
                            "type":     "image",
                            "filename": os.path.basename(str(part.partname)),
                            "ct":       getattr(part, "content_type", "image/png"),
                            "bytes":    part.blob,
                        })
                # WMF/EMF (eski Equation Editor)
                for shape in child.findall(".//v:shape", namespaces=XML_NS):
                    imagedata = shape.find(".//v:imagedata", namespaces=XML_NS)
                    if imagedata is not None:
                        rid = imagedata.get(f'{{{XML_NS["r"]}}}id')
                        if rid and rid in self.related_parts:
                            part = self.related_parts[rid]
                            blocks.append({
                                "type":     "image",
                                "filename": os.path.basename(str(part.partname)),
                                "ct":       "image/wmf",
                                "bytes":    part.blob,
                            })

            elif tag in ('oMath', 'oMathPara'):
                latex = _omml_to_latex(child)
                if latex:
                    blocks.append({"type": "formula", "latex": latex})

        return blocks

    def _extract_cell(self, cell):
        """Katak ichidagi barcha bloklarni qaytaradi"""
        all_blocks = []
        for para in cell.paragraphs:
            b = self._extract_paragraph(para)
            if not b and para.text.strip():
                b = [{"type": "text", "value": para.text.strip()}]
            all_blocks.extend(b)

        text_parts  = []
        html_parts  = []
        images      = []
        has_formula = False

        for blk in all_blocks:
            if blk["type"] == "text":
                text_parts.append(blk["value"])
                html_parts.append(escape(blk["value"]))
            elif blk["type"] == "formula":
                has_formula = True
                latex = blk["latex"]
                text_parts.append(latex)
                html_parts.append(f'\\({latex}\\)')
            elif blk["type"] == "image":
                images.append(blk)
                html_parts.append("[rasm]")
                text_parts.append("[rasm]")

        return {
            "text":        " ".join(text_parts).strip(),
            "html":        " ".join(html_parts).strip(),
            "has_formula": has_formula,
            "images":      images,
            "is_empty":    not text_parts and not images,
        }

    def parse(self):
        """Barcha jadvallardan savollar"""
        questions = []

        for table in self.doc.tables:
            if len(table.rows) < 2 or len(table.columns) < 3:
                continue

            for ri, row in enumerate(table.rows):
                cells = row.cells

                # Sarlavha qatorini o'tkazib yuborish
                first_cell_text = cells[0].text.strip().lower() if cells else ''
                if any(h in first_cell_text for h in ('№', 'no', '#', 't/r')):
                    if ri == 0 or any(c.text.strip().lower() in ('savol', 'test savoli', 'question', "to'g'ri javob") for c in cells[:3]):
                        continue
                if not any(c.text.strip() for c in cells):
                    continue

                # Ustun mapping: №[0], Savol[1], To'g'ri[2], Muqobil1[3], Muqobil2[4], Muqobil3[5]
                if len(cells) < 3:
                    continue

                q_data       = self._extract_cell(cells[1])
                correct_data = self._extract_cell(cells[2])
                wrongs       = []
                for i in range(3, min(7, len(cells))):
                    w = self._extract_cell(cells[i])
                    if not w['is_empty']:
                        wrongs.append(w)

                if q_data['is_empty']:
                    continue

                # TTS uchun mos: formulasiz va rasmsiz
                accessible = (
                    not q_data['has_formula']
                    and not q_data['images']
                    and not correct_data['has_formula']
                )

                questions.append({
                    "q":          q_data,
                    "correct":    correct_data,
                    "wrong":      wrongs,
                    "accessible": accessible,
                })

        return questions


def import_docx_to_assignment(assignment, docx_path):
    """DOCX → Question modellariga import"""
    from assessment.models import Question

    if not HAS_LXML:
        return 0, ["lxml kutubxonasi yo'q: pip install lxml"]

    try:
        parser    = AvtoBaholashDocxParser(docx_path)
        questions = parser.parse()
    except Exception as e:
        return 0, [f"Parse xatosi: {e}"]

    if not questions:
        return 0, ["Savollar topilmadi. Word jadvalda 6 ustun bo'lishi kerak: №, Savol, To'g'ri, Muqobil 1-3"]

    created = 0
    errors  = []

    for idx, item in enumerate(questions, 1):
        try:
            q       = item['q']
            correct = item['correct']
            wrongs  = item['wrong']

            # Variantlar — har biri (text, html, image)
            options = [{
                'text':  correct['text']  or '[rasm]',
                'html':  correct['html']  or '[rasm]',
                'image': correct['images'][0] if correct['images'] else None,
                'is_correct': True,
            }]
            for w in wrongs:
                options.append({
                    'text':  w['text']  or '[rasm]',
                    'html':  w['html']  or '[rasm]',
                    'image': w['images'][0] if w['images'] else None,
                    'is_correct': False,
                })

            # ARALASHTIRISH (har talabaga emas — bir marta jadvalga yozish vaqtida)
            random.shuffle(options)

            # 4 ta ustun bo'lishi shart, kerak bo'lsa bo'sh joy
            while len(options) < 4:
                options.append({'text': '', 'html': '', 'image': None, 'is_correct': False})

            # To'g'ri javob harfi
            letters = ['A', 'B', 'C', 'D']
            correct_letter = 'A'
            for li, opt in enumerate(options[:4]):
                if opt['is_correct']:
                    correct_letter = letters[li]
                    break

            # Savol matni — formulali bo'lsa HTML, oddiy bo'lsa matn
            q_text = q['html'] if q['has_formula'] else q['text']
            if not q_text or q_text == '[rasm]':
                q_text = f"{idx}-savol"

            obj = Question(
                assignment    = assignment,
                text          = q_text[:2000],
                correct_answer= correct_letter,
                option_a      = (options[0]['html'] if options[0].get('text') else '')[:500],
                option_b      = (options[1]['html'] if options[1].get('text') else '')[:500],
                option_c      = (options[2]['html'] if options[2].get('text') else '')[:500],
                option_d      = (options[3]['html'] if options[3].get('text') else '')[:500],
                order         = assignment.questions.count() + 1,
                is_accessible = item['accessible'],
            )

            # Savol rasmi
            if q['images']:
                img = q['images'][0]
                ext = os.path.splitext(img['filename'])[1] or '.png'
                obj.image.save(
                    f"q_{assignment.pk}_{idx}{ext}",
                    ContentFile(img['bytes']),
                    save=False,
                )

            # Variant rasmlari (A, B, C, D)
            for i, letter in enumerate(['A', 'B', 'C', 'D']):
                if i < len(options) and options[i].get('image'):
                    img = options[i]['image']
                    ext = os.path.splitext(img['filename'])[1] or '.png'
                    field = getattr(obj, f'image_{letter.lower()}')
                    field.save(
                        f"q_{assignment.pk}_{idx}_{letter}{ext}",
                        ContentFile(img['bytes']),
                        save=False,
                    )

            obj.save()
            created += 1
        except Exception as e:
            errors.append(f"{idx}-savol: {str(e)[:80]}")

    return created, errors
