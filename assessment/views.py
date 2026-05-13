from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.conf import settings
import json

from assessment.models import Assignment, Submission, Notification
from assessment.services.export_service import (
    assignment_pdf_response, export_gradebook_xlsx, send_telegram_message
)


@login_required
def download_pdf(request, pk):
    """Topshiriqni PDF holatida yuklab olish"""
    assignment = get_object_or_404(Assignment, pk=pk)
    # Faqat o'qituvchi yoki talabaning guruhi uchun
    if not (request.user.is_teacher or request.user.is_admin or
            (request.user.is_student and request.user.group and
             assignment.groups.filter(pk=request.user.group.pk).exists())):
        return HttpResponse("Ruxsat yo'q", status=403)
    return assignment_pdf_response(assignment)


@login_required
def export_xlsx(request, pk):
    """Baholash jurnalini Excel formatida yuklab olish"""
    assignment = get_object_or_404(Assignment, pk=pk, teacher=request.user)
    submissions = Submission.objects.filter(
        assignment=assignment
    ).select_related('student', 'student__group').order_by('student__last_name')

    xlsx_bytes = export_gradebook_xlsx(assignment, submissions)
    if not xlsx_bytes:
        return HttpResponse("Excel export xatosi", status=500)

    response = HttpResponse(
        xlsx_bytes,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    safe_name = assignment.title[:30].replace(' ', '_')
    response['Content-Disposition'] = f'attachment; filename="jurnal_{safe_name}.xlsx"'
    return response


@login_required
def download_test_template(request):
    """Test savollari Word shablonini yuklab olish"""
    return _word_template_response('test')


@login_required
def download_written_template(request):
    """Yozma ish savollari Word shablonini yuklab olish"""
    return _word_template_response('written')


def _word_template_response(template_type):
    """Word shablon yaratib qaytaradi"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        if template_type == 'test':
            # Title
            title = doc.add_heading('AvtoBaholash — Test Savollari Shabloni', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Instructions
            doc.add_paragraph()
            p = doc.add_paragraph('📌 TO\'LDIRISH KO\'RSATMALARI:')
            p.runs[0].bold = True

            instructions = [
                '• 1-satr — SARLAVHA (o\'zgartirmang!)',
                '• tr — Tartib raqami (1, 2, 3, ...)',
                '• savol — Savol matni (matn, rasm yoki formula bo\'lishi mumkin)',
                '• togri_javob — To\'g\'ri javob (A, B, C yoki D harfi)',
                '• muqobil1 — A variant matni',
                '• muqobil2 — B variant matni',
                '• muqobil3 — C variant matni',
                '• muqobil4 — D variant matni (ixtiyoriy)',
                '⚠️  Jadvalda faqat 1 ta jadval bo\'lishi kerak!',
            ]
            for inst in instructions:
                doc.add_paragraph(inst)

            doc.add_paragraph()

            # Table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'

            # Header
            headers = ['tr', 'savol', 'togri_javob', 'muqobil1', 'muqobil2', 'muqobil3']
            hdr_row = table.rows[0]
            for i, h in enumerate(headers):
                cell = hdr_row.cells[i]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.bold = True

            # Sample rows
            samples = [
                ('1', 'Ikki sonning yig\'indisi 10 bo\'lsa, biri 4 bo\'lsa, ikkinchisi?', 'A', '6', '4', '8'),
                ('2', 'O\'zbekiston poytaxti qaysi shahar?', 'A', 'Toshkent', 'Samarqand', 'Buxoro'),
            ]
            for s in samples:
                row = table.add_row()
                for i, val in enumerate(s):
                    row.cells[i].text = val

            # Empty rows
            for i in range(3, 16):
                row = table.add_row()
                row.cells[0].text = str(i)

        else:  # written
            title = doc.add_heading('AvtoBaholash — Yozma Ish Savollari Shabloni', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
            p = doc.add_paragraph('📌 TO\'LDIRISH KO\'RSATMALARI:')
            p.runs[0].bold = True

            instructions = [
                '• Har bir guruh uchun ALOHIDA jadval yarating',
                '• 1-satr, 2-ustun — Guruh nomi (masalan: "Integral hisob", "1-modul")',
                '• 2-satrdan boshlab, 2-ustun — Savol matni',
                '• 1-ustun — Tartib raqami (tizim e\'tiborga olmaydi, faqat ko\'rgazma)',
                '• Savollarga rasm yoki formula ham qo\'shish mumkin',
            ]
            for inst in instructions:
                doc.add_paragraph(inst)

            doc.add_paragraph()

            # Sample group table
            for group_num in range(1, 3):
                heading = doc.add_heading(f'{group_num}-guruh namunasi', level=2)
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'

                # Group name row
                hdr = table.rows[0]
                hdr.cells[0].text = f'{group_num}'
                hdr.cells[0].paragraphs[0].runs[0].bold = True
                hdr.cells[1].text = f'Guruh nomi shu yerga — {group_num}-guruh'
                hdr.cells[1].paragraphs[0].runs[0].bold = True

                # Sample questions
                for q_num in range(1, 4):
                    row = table.add_row()
                    row.cells[0].text = str(q_num)
                    row.cells[1].text = f'Savol {q_num} shu yerga yoziladi...'

                # Empty rows
                for q_num in range(4, 9):
                    row = table.add_row()
                    row.cells[0].text = str(q_num)

                doc.add_paragraph()

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        fname = f"{'test' if template_type == 'test' else 'yozma_ish'}_shablon.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{fname}"'
        return response

    except Exception as e:
        return HttpResponse(f"Shablon yaratishda xato: {e}", status=500)


@csrf_exempt
@require_POST
def telegram_webhook(request):
    """
    Telegram bot webhook endpoint.
    Bot sozlamalari: /setwebhook -> https://yourdomain.com/assessment/webhook/telegram/
    """
    try:
        data = json.loads(request.body)
        message = data.get('message', {})
        chat_id = message.get('chat', {}).get('id')
        text = message.get('text', '').strip()

        if not chat_id:
            return JsonResponse({'ok': True})

        bot_token = getattr(settings, 'TELEGRAM_BOT_TOKEN', '')

        if text == '/start':
            reply = (
                "👋 AvtoBaholash botiga xush kelibsiz!\n\n"
                "Bu bot orqali quyidagilarni olasiz:\n"
                "• 📚 Yangi topshiriqlar haqida xabar\n"
                "• ✅ Baho tayyor bo'lganda bildirishnoma\n"
                "• ⏰ Muddat eslatmalari\n\n"
                f"🔗 Tizimga kirish: {request.build_absolute_uri('/')}\n"
                f"🆔 Sizning chat ID: <code>{chat_id}</code>\n\n"
                "Chat ID ni AvtoBaholash profilingizga kiriting."
            )
        elif text == '/help':
            reply = (
                "📋 AvtoBaholash Bot Buyruqlar:\n\n"
                "/start - Botni ishga tushirish\n"
                "/help - Yordam\n"
                "/status - Mening holat\n\n"
                "Profilingizda chat ID ni ko'rsating."
            )
        else:
            reply = f"✅ Xabar qabul qilindi. Chat ID: <code>{chat_id}</code>"

        if bot_token:
            send_telegram_message(bot_token, chat_id, reply, parse_mode='HTML')

        return JsonResponse({'ok': True})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)})
