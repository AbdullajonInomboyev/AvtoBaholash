from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from django.conf import settings
from django.http import JsonResponse, HttpResponse
from django.db.models import Avg, Count, Q
from datetime import timedelta
import json
import logging
logger = logging.getLogger(__name__)

from accounts.models import User
from core.models import Department, Group, Subject, SubjectTeacher, ExamSchedule
from assessment.models import (Assignment, Question, Submission,
                                Feedback, Notification, AIAnalysisLog, Syllabus,
                                QuestionBank, BankQuestion, TopicQuota)


# ─── decorator ────────────────────────────────────────────────
def role_required(*roles):
    def decorator(fn):
        @login_required
        def wrapper(request, *args, **kwargs):
            if request.user.role not in roles:
                messages.error(request, "Bu sahifaga kirish huquqingiz yo'q.")
                return redirect(request.user.get_dashboard_url())
            return fn(request, *args, **kwargs)
        wrapper.__name__ = fn.__name__
        return wrapper
    return decorator


# ─── shared context ───────────────────────────────────────────
def _base_ctx(request):
    unread = Notification.objects.filter(
        recipient=request.user, is_read=False).count()
    return {'unread_count': unread}


# ══════════════════════════════════════════════════════════════
# ADMIN
# ══════════════════════════════════════════════════════════════
@login_required
def dashboard_router(request):
    """Login qilgan foydalanuvchini o'z dashboard'iga yo'naltiradi"""
    return redirect(request.user.get_dashboard_url())


@role_required('admin')
def admin_dashboard(request):
    ctx = {**_base_ctx(request),
        'departments':      Department.objects.annotate(tc=Count('staff', filter=Q(staff__role='oqituvchi')),sc=Count('groups__students'),sub_c=Count('subjects')),
        'users_total':      User.objects.count(),
        'teachers_total':   User.objects.filter(role='oqituvchi').count(),
        'students_total':   User.objects.filter(role='talaba').count(),
        'subjects_total':   Subject.objects.count(),
        'assignments_total':Assignment.objects.count(),
        'recent_users':     User.objects.order_by('-date_joined')[:8],
        'stat_cards': [
            ('Foydalanuvchilar', User.objects.count(), 'group', 'bg-blue-100 text-blue-600'),
            ("O'qituvchilar", User.objects.filter(role='oqituvchi').count(), 'school', 'bg-green-100 text-green-600'),
            ('Talabalar', User.objects.filter(role='talaba').count(), 'person', 'bg-orange-100 text-orange-600'),
            ('Fanlar', Subject.objects.count(), 'book', 'bg-purple-100 text-purple-600'),
            ('Topshiriqlar', Assignment.objects.count(), 'assignment', 'bg-yellow-100 text-yellow-600'),
        ],
    }
    return render(request, 'admin_panel/dashboard.html', ctx)


@role_required('admin')
def admin_departments(request):
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create':
            Department.objects.get_or_create(
                code=request.POST['code'],
                defaults={'name': request.POST['name'], 'description': request.POST.get('description','')})
            messages.success(request, "Kafedra qo'shildi.")
        elif act == 'delete':
            Department.objects.filter(pk=request.POST['id']).delete()
            messages.success(request, "O'chirildi.")
        return redirect('core:admin_departments')
    depts = Department.objects.annotate(
        tc=Count('staff', filter=Q(staff__role='oqituvchi')),
        sc=Count('groups__students'),
        sub_c=Count('subjects')).order_by('name')
    return render(request, 'admin_panel/departments.html', {**_base_ctx(request), 'departments': depts})



@role_required('admin')
def admin_department_detail(request, pk):
    """Kafedra ichida — o'qituvchilar, fanlar, guruhlar boshqaruvi"""
    dept = get_object_or_404(Department, pk=pk)

    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'add_teacher':
            tid = request.POST.get('teacher_id')
            if tid:
                t = User.objects.filter(pk=tid, role='oqituvchi').first()
                if t:
                    t.department = dept
                    t.save()
                    messages.success(request, f"{t.full_name} kafedraga qo'shildi.")
        elif act == 'create_teacher':
            try:
                u = User.objects.create_user(
                    username   = request.POST['username'],
                    password   = request.POST.get('password') or settings.DEFAULT_STUDENT_PASSWORD,
                    first_name = request.POST.get('first_name', ''),
                    last_name  = request.POST.get('last_name', ''),
                    email      = request.POST.get('email', ''),
                    role       = 'oqituvchi',
                    department = dept,
                )
                messages.success(request, f"{u.full_name} yaratildi va kafedraga qo'shildi.")
            except Exception as e:
                messages.error(request, f"Xato: {e}")
        elif act == 'remove_teacher':
            tid = request.POST.get('teacher_id')
            User.objects.filter(pk=tid, role='oqituvchi', department=dept).update(department=None)
            messages.success(request, "O'qituvchi olib tashlandi.")
        elif act == 'add_subject':
            try:
                Subject.objects.create(
                    name        = request.POST['name'],
                    code        = request.POST['code'],
                    department  = dept,
                    credits     = int(request.POST.get('credits', 3)),
                    semester    = int(request.POST.get('semester', 1)),
                )
                messages.success(request, "Fan qo'shildi.")
            except Exception as e:
                messages.error(request, f"Xato: {e}")
        elif act == 'remove_subject':
            Subject.objects.filter(pk=request.POST['subject_id'], department=dept).delete()
            messages.success(request, "Fan o'chirildi.")
        elif act == 'add_group':
            try:
                Group.objects.create(
                    name        = request.POST['name'],
                    year        = int(request.POST.get('year', 1)),
                    department  = dept,
                )
                messages.success(request, "Guruh qo'shildi.")
            except Exception as e:
                messages.error(request, f"Xato: {e}")
        elif act == 'remove_group':
            Group.objects.filter(pk=request.POST['group_id'], department=dept).delete()
            messages.success(request, "Guruh o'chirildi.")
        return redirect('core:admin_department_detail', pk=dept.pk)

    return render(request, 'admin_panel/department_detail.html', {
        **_base_ctx(request),
        'dept':              dept,
        'teachers':          User.objects.filter(department=dept, role='oqituvchi').order_by('last_name'),
        'subjects':          Subject.objects.filter(department=dept).order_by('name'),
        'groups':            Group.objects.filter(department=dept).annotate(sc=Count('students')).order_by('name'),
        'free_teachers':     User.objects.filter(role='oqituvchi', department__isnull=True),
    })


@role_required('admin')
def admin_users(request):
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create':
            u = User.objects.create_user(
                username=request.POST['username'],
                password=request.POST.get('password',settings.DEFAULT_STUDENT_PASSWORD),
                first_name=request.POST.get('first_name',''),
                last_name=request.POST.get('last_name',''),
                email=request.POST.get('email',''),
                role=request.POST.get('role','talaba'),
                phone=request.POST.get('phone',''))
            if request.POST.get('department'): u.department_id=request.POST['department']
            if request.POST.get('group'):      u.group_id=request.POST['group']
            u.save()
            messages.success(request, f"'{u.username}' yaratildi.")
        elif act == 'delete':
            User.objects.filter(pk=request.POST['id']).exclude(pk=request.user.pk).delete()
            messages.success(request, "O'chirildi.")
        elif act == 'toggle':
            u = get_object_or_404(User, pk=request.POST['id'])
            u.is_active = not u.is_active; u.save()
        return redirect('core:admin_users')
    role_f = request.GET.get('role','')
    qs = User.objects.select_related('department','group')
    if role_f: qs = qs.filter(role=role_f)
    from accounts.models import User as U
    return render(request, 'admin_panel/users.html', {
        **_base_ctx(request),
        'users': qs.order_by('-date_joined'),
        'departments': Department.objects.all(),
        'groups': Group.objects.all(),
        'role_filter': role_f,
        'roles': U.ROLES,
    })


@role_required('admin')
def admin_subjects(request):
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create':
            Subject.objects.create(
                name=request.POST['name'], code=request.POST['code'],
                department_id=request.POST['department'],
                credits=int(request.POST.get('credits',3)),
                semester=int(request.POST.get('semester',1)))
            messages.success(request, "Fan qo'shildi.")
        elif act == 'delete':
            Subject.objects.filter(pk=request.POST['id']).delete()
            messages.success(request, "O'chirildi.")
        return redirect('core:admin_subjects')
    return render(request, 'admin_panel/subjects.html', {
        **_base_ctx(request),
        'subjects': Subject.objects.select_related('department').order_by('name'),
        'departments': Department.objects.all(),
    })


@role_required('admin')
def admin_reports(request):
    ctx = {**_base_ctx(request),
        'total_assignments': Assignment.objects.count(),
        'total_submissions': Submission.objects.count(),
        'graded':            Submission.objects.filter(status='graded').count(),
        'avg_score':         Submission.objects.filter(final_score__isnull=False).aggregate(a=Avg('final_score'))['a'],
        'departments':       Department.objects.annotate(sub=Count('subjects'), tc=Count('staff',filter=Q(staff__role='oqituvchi'))),
    }
    return render(request, 'admin_panel/reports.html', ctx)


@role_required('admin')
def admin_settings(request):
    from django.conf import settings as django_settings
    if request.method == 'POST':
        messages.success(request, 'Sozlamalar saqlandi.')
    ctx = {**_base_ctx(request), 'settings': django_settings}
    return render(request, 'admin_panel/settings.html', ctx)


# ══════════════════════════════════════════════════════════════
# KAFEDRA MUDIRI
# ══════════════════════════════════════════════════════════════
@role_required('kafedra_mudiri')
def kafedra_dashboard(request):
    dept = request.user.department
    if not dept:
        messages.warning(request, "Sizga kafedra biriktirilmagan. Admin bilan bog'laning.")
        return render(request, 'kafedra/dashboard.html', _base_ctx(request))

    teachers  = User.objects.filter(role='oqituvchi', department=dept)
    week_ago  = timezone.now() - timedelta(days=7)
    ai_logs   = AIAnalysisLog.objects.filter(
        assignment__subject__department=dept).select_related('assignment__teacher').order_by('-created_at')[:6]

    ctx = {**_base_ctx(request),
        'dept': dept,
        'teachers_count':  teachers.count(),
        'students_count':  User.objects.filter(role='talaba', group__department=dept).count(),
        'subjects_count':  dept.subjects.count(),
        'exams_count':     ExamSchedule.objects.filter(department=dept, exam_date__gte=timezone.now()).count(),
        'teachers':        teachers.annotate(ac=Count('created_assignments'))[:8],
        'upcoming_exams':  ExamSchedule.objects.filter(department=dept, exam_date__gte=timezone.now()).order_by('exam_date')[:5],
        'ai_logs':         ai_logs,
        'weekly_new':      Assignment.objects.filter(subject__department=dept, created_at__gte=week_ago).count(),
        'weekly_subs':     Submission.objects.filter(assignment__subject__department=dept, submitted_at__gte=week_ago).count(),
        'active_students': User.objects.filter(role='talaba', group__department=dept, is_active=True).count(),
    }
    return render(request, 'kafedra/dashboard.html', ctx)


@role_required('kafedra_mudiri')
def kafedra_teachers(request):
    dept = request.user.department
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create':
            u = User.objects.create_user(
                username=request.POST['username'],
                password=request.POST.get('password',settings.DEFAULT_STUDENT_PASSWORD),
                first_name=request.POST.get('first_name',''),
                last_name=request.POST.get('last_name',''),
                email=request.POST.get('email',''),
                role='oqituvchi', department=dept,
                phone=request.POST.get('phone',''))
            messages.success(request, f"'{u.full_name}' qo'shildi.")
        elif act == 'assign':
            st, _ = SubjectTeacher.objects.get_or_create(
                teacher_id=request.POST['teacher_id'],
                subject_id=request.POST['subject_id'])
            st.groups.set(request.POST.getlist('groups'))
            messages.success(request, "Fan biriktirildi.")
        elif act == 'delete':
            teacher_to_del = User.objects.filter(pk=request.POST['id'], role='oqituvchi', department=dept).first()
            if teacher_to_del:
                a_count = teacher_to_del.created_assignments.count()
                teacher_to_del.delete()
                if a_count:
                    messages.warning(request, f"O'qituvchi o'chirildi. Uning {a_count} ta topshirig'i ham o'chirildi.")
                else:
                    messages.success(request, "O'chirildi.")
        return redirect('core:kafedra_teachers')
    teachers = User.objects.filter(role='oqituvchi', department=dept).select_related('department').annotate(
        sc=Count('teaching_subjects'), ac=Count('created_assignments'))
    return render(request, 'kafedra/teachers.html', {
        **_base_ctx(request),
        'teachers': teachers, 'dept': dept,
        'subjects': Subject.objects.filter(department=dept),
        'groups':   Group.objects.filter(department=dept),
        'subject_teachers': SubjectTeacher.objects.filter(subject__department=dept).select_related('teacher','subject'),
    })


@role_required('kafedra_mudiri')
def kafedra_students(request):
    dept = request.user.department
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create':
            u = User.objects.create_user(
                username=request.POST['username'],
                password=request.POST.get('password', settings.DEFAULT_STUDENT_PASSWORD),
                first_name=request.POST.get('first_name', ''),
                last_name=request.POST.get('last_name', ''),
                email=request.POST.get('email', ''),
                role='talaba',
                student_id=request.POST.get('student_id', ''),
                group_id=request.POST.get('group') or None,
                is_accessible='is_accessible' in request.POST,
            )
            if u.is_accessible:
                messages.success(request, f"'{u.full_name}' qo'shildi. ♿ Inklyuziv ta'lim rejimi yoqilgan.")
            else:
                messages.success(request, f"'{u.full_name}' qo'shildi.")
        elif act == 'delete':
            User.objects.filter(pk=request.POST['id'], role='talaba', group__department=dept).delete()
            messages.success(request, "O'chirildi.")
        return redirect('core:kafedra_students')
    search            = request.GET.get('q', '').strip()
    group_filter      = request.GET.get('group', '')
    accessible_filter = request.GET.get('accessible', '')

    students = User.objects.filter(
        role='talaba', group__department=dept
    ).select_related('group').annotate(sub_c=Count('submissions')).order_by('last_name', 'first_name')

    if group_filter:
        students = students.filter(group_id=group_filter)
    if accessible_filter == '1':
        students = students.filter(is_accessible=True)
    elif accessible_filter == '0':
        students = students.filter(is_accessible=False)
    if search:
        students = students.filter(
            Q(first_name__icontains=search) | Q(last_name__icontains=search) |
            Q(username__icontains=search)   | Q(student_id__icontains=search)
        )

    return render(request, 'kafedra/students.html', {
        **_base_ctx(request),
        'students':          students,
        'dept':              dept,
        'groups':            Group.objects.filter(department=dept),
        'search':            search,
        'group_filter':      group_filter,
        'accessible_filter': accessible_filter,
    })


@role_required('kafedra_mudiri')
def kafedra_subjects(request):
    dept = request.user.department
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create_subject':
            Subject.objects.create(
                name=request.POST['name'], code=request.POST['code'],
                department=dept, credits=int(request.POST.get('credits',3)),
                semester=int(request.POST.get('semester',1)))
            messages.success(request, "Fan qo'shildi.")
        elif act == 'create_group':
            Group.objects.create(name=request.POST['group_name'], department=dept,
                                 year=int(request.POST.get('year',1)))
            messages.success(request, "Guruh qo'shildi.")
        elif act == 'delete_subject':
            Subject.objects.filter(pk=request.POST['id'], department=dept).delete()
            messages.success(request, "O'chirildi.")
        elif act == 'delete_group':
            Group.objects.filter(pk=request.POST['id'], department=dept).delete()
            messages.success(request, "Guruh o'chirildi.")
        elif act == 'assign_teacher':
            try:
                subject = Subject.objects.get(pk=request.POST['subject_id'], department=dept)
                teacher = User.objects.get(pk=request.POST['teacher_id'], role='oqituvchi', department=dept)
                st, created = SubjectTeacher.objects.get_or_create(
                    subject=subject, teacher=teacher
                )
                group_ids = request.POST.getlist('groups')
                if group_ids:
                    st.groups.set(group_ids)
                if created:
                    messages.success(request, f"{teacher.full_name} → {subject.name} fanida biriktirildi")
                else:
                    messages.info(request, "Bu o'qituvchi allaqachon biriktirilgan, guruhlar yangilandi")
            except Exception as e:
                messages.error(request, f"Xato: {e}")
        elif act == 'unassign_teacher':
            try:
                st = SubjectTeacher.objects.get(pk=request.POST['st_id'], subject__department=dept)
                st.delete()
                messages.success(request, "O'qituvchi olib tashlandi")
            except Exception as e:
                messages.error(request, f"Xato: {e}")
        return redirect('core:kafedra_subjects')
    return render(request, 'kafedra/subjects.html', {
        **_base_ctx(request), 'dept': dept,
        'subjects':         Subject.objects.filter(department=dept).annotate(tc=Count('subject_teachers')),
        'groups':           Group.objects.filter(department=dept).annotate(sc=Count('students')),
        'teachers':         User.objects.filter(role='oqituvchi', department=dept),
        'subject_teachers': SubjectTeacher.objects.filter(subject__department=dept).select_related('teacher','subject'),
    })


@role_required('kafedra_mudiri')
def kafedra_exam_schedule(request):
    dept = request.user.department
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create':
            ex = ExamSchedule.objects.create(
                subject_id=request.POST['subject'], department=dept,
                exam_date=request.POST['exam_date'],
                room=request.POST.get('room',''),
                duration_minutes=int(request.POST.get('duration',120)),
                notes=request.POST.get('notes',''),
                created_by=request.user)
            ex.groups.set(request.POST.getlist('groups'))
            messages.success(request, "Imtihon qo'shildi.")
        elif act == 'delete':
            ExamSchedule.objects.filter(pk=request.POST['id'], department=dept).delete()
            messages.success(request, "O'chirildi.")
        return redirect('core:kafedra_exam_schedule')
    return render(request, 'kafedra/exam_schedule.html', {
        **_base_ctx(request), 'dept': dept,
        'exams':    ExamSchedule.objects.filter(department=dept).prefetch_related('groups').order_by('exam_date'),
        'subjects': Subject.objects.filter(department=dept),
        'groups':   Group.objects.filter(department=dept),
        'now':      timezone.now(),
    })



@role_required('kafedra_mudiri')
def kafedra_exam_detail(request, pk):
    """Imtihon ichi — topshiriqlarni biriktirish/uzish"""
    dept = request.user.department
    exam = get_object_or_404(ExamSchedule, pk=pk, department=dept)

    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'attach_assignment':
            aid = request.POST.get('assignment_id')
            a = Assignment.objects.filter(pk=aid, subject=exam.subject).first()
            if a:
                exam.assignments.add(a)
                messages.success(request, f"'{a.title}' biriktirildi.")
            else:
                messages.error(request, "Topshiriq topilmadi yoki boshqa fanga tegishli.")
        elif act == 'detach_assignment':
            aid = request.POST.get('assignment_id')
            exam.assignments.remove(aid)
            messages.success(request, "Topshiriq uzildi.")
        elif act == 'update':
            try:
                exam.exam_date        = request.POST.get('exam_date') or exam.exam_date
                exam.room             = request.POST.get('room', exam.room)
                exam.duration_minutes = int(request.POST.get('duration', exam.duration_minutes))
                exam.notes            = request.POST.get('notes', exam.notes)
                exam.save()
                messages.success(request, "Imtihon yangilandi.")
            except Exception as e:
                messages.error(request, f"Xato: {e}")
        return redirect('core:kafedra_exam_detail', pk=exam.pk)

    # Shu fan bo'yicha mavjud topshiriqlar (hali biriktirilmaganlar)
    attached_ids = list(exam.assignments.values_list('pk', flat=True))
    available_assignments = Assignment.objects.filter(
        subject=exam.subject
    ).exclude(pk__in=attached_ids).select_related('teacher').order_by('-created_at')

    return render(request, 'kafedra/exam_detail.html', {
        **_base_ctx(request),
        'dept':                  dept,
        'exam':                  exam,
        'attached_assignments':  exam.assignments.select_related('teacher').prefetch_related('questions').all(),
        'available_assignments': available_assignments,
    })


@role_required('kafedra_mudiri')
def kafedra_ai_journal(request):
    dept = request.user.department
    teacher_id = request.GET.get('teacher')
    qs = Assignment.objects.filter(subject__department=dept).select_related('teacher','subject').order_by('-created_at')
    if teacher_id:
        qs = qs.filter(teacher_id=teacher_id)
    rows = []
    for a in qs[:60]:
        rows.append({'assignment': a, 'log': a.ai_logs.order_by('-created_at').first()})
    return render(request, 'kafedra/ai_journal.html', {
        **_base_ctx(request), 'dept': dept, 'rows': rows,
        'teachers': User.objects.filter(role='oqituvchi', department=dept),
        'sel_teacher': teacher_id,
    })


# ══════════════════════════════════════════════════════════════
# O'QITUVCHI
# ══════════════════════════════════════════════════════════════
@role_required('oqituvchi')
def teacher_dashboard(request):
    t = request.user
    my_sts = SubjectTeacher.objects.filter(teacher=t).select_related('subject')
    pending_grade   = Submission.objects.filter(assignment__teacher=t, status='submitted').count()
    pending_feedback= Feedback.objects.filter(submission__assignment__teacher=t, status='open').count()
    ctx = {**_base_ctx(request),
        'subjects_count':  my_sts.count(),
        'assignments_count': Assignment.objects.filter(teacher=t).count(),
        'pending_grade':   pending_grade,
        'pending_feedback': pending_feedback,
        'my_subjects':     my_sts[:6],
        'active_assignments': Assignment.objects.filter(teacher=t, status='active', deadline__gte=timezone.now()).order_by('deadline')[:5],
        'recent_submissions': Submission.objects.filter(assignment__teacher=t).select_related('student','assignment').order_by('-submitted_at')[:6],
    }
    return render(request, 'teacher/dashboard.html', ctx)


@role_required('oqituvchi')
def teacher_subjects(request):
    my_sts = SubjectTeacher.objects.filter(teacher=request.user).select_related('subject__department').prefetch_related('groups')
    return render(request, 'teacher/subjects.html', {**_base_ctx(request), 'my_subjects': my_sts})


@role_required('oqituvchi')
def teacher_assignments(request):
    t   = request.user
    sid = request.GET.get('subject')

    # O'chirish so'rovi
    if request.method == 'POST' and request.POST.get('action') == 'delete':
        aid = request.POST.get('assignment_id')
        deleted = Assignment.objects.filter(pk=aid, teacher=t).delete()
        if deleted[0]:
            messages.success(request, "Topshiriq o'chirildi.")
        return redirect('core:teacher_assignments')

    qs  = Assignment.objects.filter(teacher=t).annotate(
        sub_c=Count('submissions'),
        graded_c=Count('submissions', filter=Q(submissions__status='graded'))).order_by('-created_at')
    if sid: qs = qs.filter(subject_id=sid)
    return render(request, 'teacher/assignments.html', {
        **_base_ctx(request),
        'assignments': qs,
        'my_subjects': SubjectTeacher.objects.filter(teacher=t).select_related('subject'),
        'sel_subject': sid,
    })


@role_required('oqituvchi')
def teacher_create_assignment(request):
    t    = request.user
    my_sts = SubjectTeacher.objects.filter(teacher=t).select_related('subject').prefetch_related('groups')

    # URL ?subject=<pk> orqali kelgan bo'lsa avtomatik tanlash
    preselected_subject = request.GET.get('subject', '')

    if request.method == 'POST':
        a = Assignment.objects.create(
            title           = request.POST['title'],
            assignment_type = request.POST['assignment_type'],
            subject_id      = request.POST['subject'],
            teacher         = t,
            description     = request.POST.get('description',''),
            instructions    = request.POST.get('instructions',''),
            deadline        = request.POST['deadline'],
            duration_minutes= int(request.POST.get('duration_minutes',60)),
            max_score       = float(request.POST.get('max_score',100)),
            shuffle_questions='shuffle_questions' in request.POST,
            show_review_to_student='show_review_to_student' in request.POST,
            questions_per_student=int(request.POST.get('questions_per_student',0)),
            allowed_file_types=[x.strip() for x in request.POST.get('allowed_file_types','').split(',') if x.strip()],
        )
        a.groups.set(request.POST.getlist('groups'))
        messages.success(request, f"'{a.title}' yaratildi. Endi savollar qo'shing.")
        return redirect('core:teacher_edit_assignment', pk=a.pk)
    all_groups = []
    for st in my_sts:
        all_groups.extend(st.groups.all())
    # Har bir guruhda inklyuziv talabalar soni
    from accounts.models import User as UModel
    groups_with_accessible = {}
    for g in list({g.pk: g for g in all_groups}.values()):
        acc = UModel.objects.filter(group=g, role='talaba', is_accessible=True).count()
        total = UModel.objects.filter(group=g, role='talaba').count()
        groups_with_accessible[g.pk] = {'accessible': acc, 'total': total}

    return render(request, 'teacher/create_assignment.html', {
        **_base_ctx(request),
        'my_subjects':         my_sts,
        'all_groups':          list({g.pk: g for g in all_groups}.values()),
        'groups_accessible':   groups_with_accessible,
        'preselected_subject': preselected_subject,
        'assignment_types':    [('test','Test','quiz','Variantli savollar, AI avtomatik baholaydi'),('written','Yozma ish','edit_note','Matn, rasm yoki fayl yuklash')],
    })



@login_required
def teacher_edit_question(request, pk):
    """Savolni alohida sahifada tahrirlash (LaTeX, rasm bilan)"""
    if not request.user.is_teacher:
        return redirect('core:teacher_dashboard')
    q = get_object_or_404(Question, pk=pk, assignment__teacher=request.user)

    if request.method == 'POST':
        q.text           = request.POST.get('text', q.text)
        q.option_a       = request.POST.get('option_a', q.option_a)
        q.option_b       = request.POST.get('option_b', q.option_b)
        q.option_c       = request.POST.get('option_c', q.option_c)
        q.option_d       = request.POST.get('option_d', q.option_d)
        q.correct_answer = request.POST.get('correct_answer', q.correct_answer)
        q.topic          = request.POST.get('topic', q.topic)
        q.difficulty     = request.POST.get('difficulty', q.difficulty)
        try:
            q.points     = int(request.POST.get('points', q.points) or 1)
        except ValueError:
            pass

        # Rasmlar
        if 'image' in request.FILES:   q.image   = request.FILES['image']
        if 'image_a' in request.FILES: q.image_a = request.FILES['image_a']
        if 'image_b' in request.FILES: q.image_b = request.FILES['image_b']
        if 'image_c' in request.FILES: q.image_c = request.FILES['image_c']
        if 'image_d' in request.FILES: q.image_d = request.FILES['image_d']

        # Rasm o'chirish
        for fld in ('image', 'image_a', 'image_b', 'image_c', 'image_d'):
            if request.POST.get(f'{fld}_clear') == '1':
                f = getattr(q, fld)
                if f:
                    f.delete(save=False)
                setattr(q, fld, None)

        q.save()
        messages.success(request, "Savol yangilandi.")
        return redirect('core:teacher_edit_assignment', pk=q.assignment.pk)

    return render(request, 'teacher/edit_question.html', {
        **_base_ctx(request),
        'q':          q,
        'assignment': q.assignment,
    })



@role_required('oqituvchi')
def teacher_edit_assignment(request, pk):
    a = get_object_or_404(Assignment, pk=pk, teacher=request.user)
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'add_question':
            q = Question(assignment=a, text=request.POST['text'],
                topic=request.POST.get('topic',''),
                difficulty=request.POST.get('difficulty', 'medium'),
                points=int(request.POST.get('points', 1) or 1),
                correct_answer=request.POST['correct_answer'],
                option_a=request.POST['option_a'],
                option_b=request.POST['option_b'],
                option_c=request.POST['option_c'],
                option_d=request.POST.get('option_d',''),
                order=a.questions.count()+1)
            if 'image' in request.FILES: q.image = request.FILES['image']
            q.save()
            messages.success(request, "Savol qo'shildi.")
        elif act == 'del_question':
            Question.objects.filter(pk=request.POST['qid'], assignment=a).delete()
        elif act == 'upload_docx':
            f = request.FILES.get('docx_file')
            if f:
                import tempfile, os as _os
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    for chunk in f.chunks(): tmp.write(chunk)
                    tmp_path = tmp.name
                try:
                    from assessment.services.docx_importer import import_docx_to_assignment
                    n, errs = import_docx_to_assignment(a, tmp_path)
                    if n > 0:
                        messages.success(request, f"✅ {n} ta savol import qilindi (matn, formula, rasm).")
                    else:
                        messages.warning(request, "Savollar topilmadi — jadval formatini tekshiring.")
                    for e in errs[:3]:
                        messages.warning(request, e)
                except Exception as ex:
                    messages.error(request, f"Import xatosi: {ex}")
                finally:
                    _os.unlink(tmp_path)
        elif act == 'activate':
            # Tanlangan talabalar
            student_ids = request.POST.getlist('student_ids')
            if student_ids:
                a.allowed_students.set(student_ids)
            else:
                a.allowed_students.clear()  # bo'sh bo'lsa - barchasiga
            a.status = 'active'; a.save()
            _notify_students(a)
            # AI tekshiruvini orqada ishga tushiramiz (sahifa darhol ochiladi)
            import threading
            def _ai_check_async(assignment_pk):
                try:
                    from assessment.models import Assignment as _A
                    asg = _A.objects.get(pk=assignment_pk)
                    _ai_check(asg)
                except Exception as e:
                    logger.error(f"AI check async error: {e}")
            threading.Thread(target=_ai_check_async, args=(a.pk,), daemon=True).start()
            messages.success(request, "✅ Topshiriq faollashtirildi! AI tekshiruvi orqada davom etadi.")
        elif act == 'close':
            a.status = 'closed'; a.save()
            messages.success(request, "Topshiriq yopildi.")
        elif act == 'save_info':
            a.title       = request.POST.get('title', a.title)
            a.description = request.POST.get('description', a.description)
            a.deadline    = request.POST.get('deadline', a.deadline)
            a.show_review_to_student = 'show_review_to_student' in request.POST
            a.save(); messages.success(request, "Saqlandi.")
        return redirect('core:teacher_edit_assignment', pk=pk)
    # Guruh talabalari (faollashtirish uchun)
    group_students = User.objects.filter(
        role='talaba', group__in=a.groups.all()
    ).select_related('group').order_by('group__name', 'last_name', 'first_name')
    allowed_ids = set(a.allowed_students.values_list('pk', flat=True))

    return render(request, 'teacher/edit_assignment.html', {
        **_base_ctx(request), 'assignment': a,
        'questions':       a.questions.order_by('order'),
        'syllabus':        a.subject.syllabi.order_by('-uploaded_at').first(),
        'group_students':  group_students,
        'allowed_ids':     allowed_ids,
    })



def _parse_question_from_cells(cells, order=1):
    """
    Word jadvaldan savol yaratadi.
    Ustunlar: [tr, savol, togri_javob_MATNI, noto'gri1, noto'gri2, noto'gri3]
    To'g'ri javob matn sifatida beriladi, tizim o'zi aralashtirib harfini belgilaydi.
    """
    import random
    if len(cells) < 3 or not cells[1]:
        return None

    correct_text = cells[2].strip()
    wrong_opts   = [cells[i].strip() for i in range(3, min(6, len(cells))) if cells[i].strip()]

    # Barcha variantlar
    all_opts = [correct_text] + wrong_opts
    random.shuffle(all_opts)
    while len(all_opts) < 3:
        all_opts.append('')

    # To'g'ri javob harfi
    letters = ['A','B','C','D']
    correct_letter = 'A'
    for i, opt in enumerate(all_opts[:4]):
        if opt.strip().lower() == correct_text.strip().lower():
            correct_letter = letters[i]
            break

    return dict(
        text          = cells[1],
        correct_answer= correct_letter,
        option_a      = all_opts[0] if len(all_opts) > 0 else '',
        option_b      = all_opts[1] if len(all_opts) > 1 else '',
        option_c      = all_opts[2] if len(all_opts) > 2 else '',
        option_d      = all_opts[3] if len(all_opts) > 3 else '',
        order         = order,
    )



def _import_docx_with_formulas(request, assignment, docx_file):
    """
    Formulali Word faylni import qiladi.
    LibreOffice orqali HTML ga o'tkazib, formulalarni rasm sifatida saqlaydi.
    Ustunlar: №, Savol, To'g'ri javob, Muqobil 1, Muqobil 2, Muqobil 3
    """
    import subprocess, zipfile, re, os, shutil, random, tempfile
    from django.core.files.base import ContentFile

    tmp_dir = tempfile.mkdtemp()
    try:
        # Vaqtinchalik faylga saqlash
        docx_path = os.path.join(tmp_dir, 'input.docx')
        with open(docx_path, 'wb') as f:
            for chunk in docx_file.chunks():
                f.write(chunk)

        # LibreOffice -> HTML
        r = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'html',
             '--outdir', tmp_dir, docx_path],
            capture_output=True, text=True, timeout=120
        )
        html_files = [f for f in os.listdir(tmp_dir) if f.endswith('.html')]
        if not html_files:
            return 0, ['LibreOffice konvertatsiya xatosi']

        html = open(os.path.join(tmp_dir, html_files[0]), encoding='utf-8', errors='ignore').read()
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', html, re.DOTALL)

        def extract_cell(cell, tmp_dir):
            """Katak ichidan matn va rasmlarni chiqaradi"""
            imgs  = re.findall(r'src="([^"]+\.(?:gif|png|jpg))"', cell)
            text  = re.sub(r'<[^>]+>', ' ', cell)
            text  = re.sub(r'\s+', ' ', text).strip()
            return text, [os.path.join(tmp_dir, i) for i in imgs if os.path.exists(os.path.join(tmp_dir, i))]

        created = 0
        errors  = []

        for ri, row in enumerate(rows[1:], 1):  # Sarlavhani o'tkazib
            cells = re.findall(r'<td[^>]*>(.*?)</td>', row, re.DOTALL)
            if len(cells) < 3:
                continue

            parts = [extract_cell(c, tmp_dir) for c in cells]

            # Savol [1], to'g'ri [2], noto'g'rilar [3,4,5]
            def combine(text, imgs):
                t = text.strip()
                return t, imgs[0] if imgs else None  # Birinchi rasm asosiy

            q_text,  q_img  = combine(*parts[1])  if len(parts) > 1 else ('', None)
            cor_text, cor_img = combine(*parts[2]) if len(parts) > 2 else ('', None)
            w1_text,  w1_img = combine(*parts[3])  if len(parts) > 3 else ('', None)
            w2_text,  w2_img = combine(*parts[4])  if len(parts) > 4 else ('', None)
            w3_text,  w3_img = combine(*parts[5])  if len(parts) > 5 else ('', None)

            if not q_text and not q_img:
                continue

            # Variantlarni aralashtiramiz
            all_opts  = [(cor_text, cor_img), (w1_text, w1_img), (w2_text, w2_img)]
            if w3_text or w3_img:
                all_opts.append((w3_text, w3_img))
            random.shuffle(all_opts)

            letters = ['A','B','C','D']
            correct_letter = 'A'
            for li, (t, _) in enumerate(all_opts):
                if t.strip().lower() == cor_text.strip().lower() and t.strip():
                    correct_letter = letters[li]
                    break
                if not cor_text.strip() and li == 0:
                    # To'g'ri javob ham rasm bo'lsa
                    correct_letter = 'A'
                    # Rasm asosida aralash tartib uchun first element to'g'ri
                    all_opts = [(cor_text, cor_img)] + [(w1_text,w1_img),(w2_text,w2_img)]
                    if w3_text or w3_img: all_opts.append((w3_text,w3_img))
                    break

            try:
                q = Question(
                    assignment    = assignment,
                    text          = q_text or f"{ri}-savol",
                    correct_answer= correct_letter,
                    option_a      = all_opts[0][0] if len(all_opts) > 0 else '',
                    option_b      = all_opts[1][0] if len(all_opts) > 1 else '',
                    option_c      = all_opts[2][0] if len(all_opts) > 2 else '',
                    option_d      = all_opts[3][0] if len(all_opts) > 3 else '',
                    order         = assignment.questions.count() + 1,
                )
                # Savol rasmi
                if q_img and os.path.exists(q_img):
                    with open(q_img, 'rb') as f:
                        q.image.save(
                            f"q_{assignment.pk}_{ri}.gif",
                            ContentFile(f.read()),
                            save=False
                        )
                q.save()
                created += 1
            except Exception as e:
                errors.append(f"{ri}-qator: {str(e)[:80]}")

        return created, errors

    except Exception as e:
        return 0, [f"Import xatosi: {str(e)[:100]}"]
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

def _import_docx(request, assignment):
    import tempfile, os
    f = request.FILES.get('docx_file')
    if not f: return
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        for chunk in f.chunks(): tmp.write(chunk)
        tmp_path = tmp.name
    try:
        from docx import Document
        doc = Document(tmp_path)
        n = 0
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                if i == 0: continue
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 3 and cells[1]:
                    qdata = _parse_question_from_cells(cells, assignment.questions.count()+1)
                    if qdata:
                        Question.objects.create(assignment=assignment, **qdata)
                        n += 1
        messages.success(request, f"{n} ta savol import qilindi.")
    except Exception as e:
        messages.error(request, f"Import xatosi: {e}")
    finally:
        os.unlink(tmp_path)


def _ai_check(assignment):
    from assessment.services.ai_service import check_syllabus_compliance, check_accessibility
    syllabus = assignment.subject.syllabi.order_by('-uploaded_at').first()
    topics   = syllabus.topics if syllabus else []
    q_text   = "\n".join(f"{i+1}. {q.text}" for i,q in enumerate(assignment.questions.all()[:10]))
    result   = check_syllabus_compliance(assignment.title, assignment.description, q_text, topics)
    assignment.ai_syllabus_score    = result.get('score', 0)
    assignment.ai_syllabus_feedback = result.get('feedback','')
    assignment.ai_checked_at        = timezone.now()
    assignment.save()
    AIAnalysisLog.objects.create(
        assignment=assignment, analysis_type='syllabus_check',
        result=result, score=result.get('score'), feedback=result.get('feedback',''))

    # Accessibility (TTS) tekshiruvi — ko'zi ojiz talabalar uchun
    if assignment.assignment_type == 'test':
        questions = list(assignment.questions.values('id','text','option_a','option_b','option_c','image'))
        q_data = [{'id': q['id'], 'text': q['text'],
                   'option_a': q['option_a'], 'option_b': q['option_b'],
                   'option_c': q['option_c'], 'image': q['image']} for q in questions]
        acc_result = check_accessibility(q_data)
        # Har bir savolga is_accessible belgilab qo'yamiz
        accessible_ids = set(acc_result.get('accessible_ids', []))
        for q in assignment.questions.all():
            q.is_accessible = q.pk in accessible_ids
            q.save(update_fields=['is_accessible'])
        AIAnalysisLog.objects.create(
            assignment=assignment, analysis_type='accessibility',
            result=acc_result,
            feedback=acc_result.get('feedback',''))


def _notify_students(assignment):
    from assessment.services.ai_service import send_notification
    # Agar allowed_students belgilangan bo'lsa - faqat ularga, aks holda barcha guruh talabalariga
    if assignment.allowed_students.exists():
        students = assignment.allowed_students.all()
    else:
        students = User.objects.filter(role='talaba', group__in=assignment.groups.all())
    for s in students:
        send_notification(s,
            f"Yangi topshiriq: {assignment.title}",
            f"{assignment.subject.name} fanidan yangi topshiriq. Muddat: {assignment.deadline.strftime('%d.%m.%Y %H:%M')}",
            'info', f"/assessment/submit/{assignment.pk}/")


@role_required('oqituvchi')
def teacher_upload_syllabus(request, subject_id):
    subject = get_object_or_404(Subject, pk=subject_id)
    if request.method == 'POST':
        topics = [t.strip() for t in request.POST.get('topics','').split('\n') if t.strip()]
        syl = Syllabus(subject=subject, teacher=request.user, topics=topics)
        if 'file' in request.FILES: syl.file = request.FILES['file']
        syl.save()
        messages.success(request, "Sillabus yuklandi.")
        return redirect('core:teacher_subjects')
    existing = subject.syllabi.order_by('-uploaded_at').first()
    return render(request, 'teacher/upload_syllabus.html', {
        **_base_ctx(request), 'subject': subject, 'existing': existing})


@role_required('oqituvchi')
def teacher_grade_book(request):
    t   = request.user
    sid = request.GET.get('subject')
    aid = request.GET.get('assignment')
    subjects    = SubjectTeacher.objects.filter(teacher=t).select_related('subject')
    sel_subject = sel_assignment = None
    assignments = submissions = []
    if sid:
        sel_subject = get_object_or_404(Subject, pk=sid)
        assignments = Assignment.objects.filter(teacher=t, subject=sel_subject)
    if aid:
        sel_assignment = get_object_or_404(Assignment, pk=aid, teacher=t)
        submissions    = Submission.objects.filter(assignment=sel_assignment).select_related('student','student__group').order_by('student__last_name')
    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'grade':
            sub = get_object_or_404(Submission, pk=request.POST['sub_id'], assignment__teacher=t)
            sub.final_score  = float(request.POST.get('final_score', sub.ai_score or 0))
            sub.teacher_note = request.POST.get('teacher_note','')
            sub.status       = 'graded'
            sub.graded_by    = t
            sub.graded_at    = timezone.now()
            sub.save()
            from assessment.services.ai_service import send_notification
            send_notification(sub.student, f"Baho tayyor: {sub.assignment.title}",
                f"Bahongiz: {sub.final_score:.0f}/100", 'success', f"/assessment/result/{sub.pk}/")
            messages.success(request, "Baho saqlandi.")
        elif act == 'ai_grade_all' and sel_assignment:
            _ai_grade_all(sel_assignment)
            messages.success(request, "AI baholash yakunlandi.")
        elif act == 'save_all':
            for k, v in request.POST.items():
                if k.startswith('score_'):
                    sub_pk = k.split('_')[1]
                    try:
                        sub = Submission.objects.get(pk=sub_pk, assignment__teacher=t)
                        if v.strip():
                            sub.final_score = float(v)
                            sub.teacher_note = request.POST.get(f'note_{sub_pk}', sub.teacher_note)
                            sub.status = 'graded'
                            sub.graded_by = t
                            sub.graded_at = timezone.now()
                            sub.save()
                    except Exception as e:
                        logger.warning(f"Score saqlash xato: {e}")
            messages.success(request, "Barcha baholar saqlandi.")
        elif act == 'export' and sel_assignment:
            return _export_xlsx(sel_assignment, submissions)
        from django.urls import reverse
        return redirect(reverse('core:teacher_grade_book') + f"?subject={sid or ''}&assignment={aid or ''}")
    return render(request, 'teacher/grade_book.html', {
        **_base_ctx(request),
        'subjects': subjects, 'sel_subject': sel_subject,
        'sel_assignment': sel_assignment, 'assignments': assignments,
        'submissions': submissions,
    })


def _ai_grade_all(assignment):
    from assessment.services.ai_service import grade_test, grade_written
    questions = list(Question.objects.filter(assignment=assignment))
    for sub in Submission.objects.filter(assignment=assignment, status='submitted'):
        try:
            if assignment.assignment_type == 'test':
                qwa = [{'text': q.text, 'correct_answer': q.correct_answer,
                        'student_answer': sub.test_answers.get(str(q.pk),'')} for q in questions]
                r = grade_test(assignment.title, qwa)
            elif assignment.assignment_type == 'written':
                r = grade_written(assignment.title, assignment.subject.name,
                                  assignment.description, sub.text_answer)
            else:
                r = {'score': 0, 'feedback': "Fayl topshiriq qo'lda baholanadi."}
            sub.ai_score=r.get('score',0); sub.ai_feedback=r.get('feedback','')
            sub.ai_graded_at=timezone.now(); sub.save()
        except Exception as e:
            logger.warning(f"Xato: {e}")


def _export_xlsx(assignment, submissions):
    from assessment.services.export_service import export_gradebook_xlsx
    xlsx_bytes = export_gradebook_xlsx(assignment, submissions)
    if not xlsx_bytes:
        return HttpResponse("Excel xatosi", status=500)
    resp = HttpResponse(xlsx_bytes,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    safe = assignment.title[:30].replace(' ','_')
    resp['Content-Disposition'] = f'attachment; filename="jurnal_{safe}.xlsx"'
    return resp


@role_required('oqituvchi')
def teacher_feedback(request):
    t = request.user
    fbs = Feedback.objects.filter(submission__assignment__teacher=t).select_related(
        'student','submission__assignment').order_by('-created_at')
    if request.method == 'POST':
        fb = get_object_or_404(Feedback, pk=request.POST['fb_id'], submission__assignment__teacher=t)
        fb.teacher_response = request.POST.get('response','')
        fb.status = 'answered'; fb.responded_at = timezone.now(); fb.save()
        from assessment.services.ai_service import send_notification
        send_notification(fb.student, "Shikoyatingizga javob berildi",
            f"O'qituvchi javob berdi: {fb.teacher_response[:80]}", 'info')
        messages.success(request, "Javob yuborildi.")
        return redirect('core:teacher_feedback')
    return render(request, 'teacher/feedback.html', {**_base_ctx(request), 'feedbacks': fbs})


# ══════════════════════════════════════════════════════════════
# TALABA
# ══════════════════════════════════════════════════════════════
@role_required('talaba')
def student_dashboard(request):
    s     = request.user
    group = s.group
    if not group:
        from django.contrib import messages as _msg
        _msg.warning(request, "Guruhingiz belgilanmagan. Administrator yoki kafedra mudiri bilan bog'laning.")
    # Faqat ushbu talabaga ruxsat etilgan topshiriqlar
    # (allowed_students bo'sh bo'lsa - barchaga, aks holda faqat ro'yxatdagilarga)
    qs = Assignment.objects.filter(groups=group, status='active').filter(
        Q(allowed_students__isnull=True) | Q(allowed_students=s)
    ).distinct().order_by('deadline') if group else Assignment.objects.none()
    done  = set(Submission.objects.filter(student=s).values_list('assignment_id', flat=True))
    qs    = qs.select_related('subject', 'teacher') if hasattr(qs, 'select_related') else qs
    recent_grades = Submission.objects.filter(student=s, status='graded').select_related('assignment').order_by('-graded_at')[:5]
    avg   = Submission.objects.filter(student=s, final_score__isnull=False).aggregate(a=Avg('final_score'))['a']
    notifs = Notification.objects.filter(recipient=s, is_read=False).order_by('-created_at')[:5]
    ctx = {**_base_ctx(request),
        'pending': qs.exclude(pk__in=done)[:5],
        'pending_count': qs.exclude(pk__in=done).count(),
        'completed_count': len(done),
        'avg_score': round(avg,1) if avg else None,
        'recent_grades': recent_grades,
        'notifications': notifs,
        'now': timezone.now(),
    }
    return render(request, 'student/dashboard.html', ctx)


@role_required('talaba')
def student_assignments(request):
    s     = request.user
    group = s.group
    sf    = request.GET.get('status','all')
    # Faqat ushbu talabaga ruxsat etilgan topshiriqlar
    # (allowed_students bo'sh bo'lsa - barchaga, aks holda faqat ro'yxatdagilarga)
    qs = Assignment.objects.filter(groups=group, status='active').filter(
        Q(allowed_students__isnull=True) | Q(allowed_students=s)
    ).distinct().order_by('deadline') if group else Assignment.objects.none()
    subs  = {sub.assignment_id: sub for sub in Submission.objects.filter(student=s)}
    rows  = []
    for a in qs:
        sub    = subs.get(a.pk)
        status = 'submitted' if sub else ('expired' if a.is_expired else 'pending')
        if sf != 'all' and status != sf: continue
        rows.append({'assignment': a, 'submission': sub, 'status': status})
    return render(request, 'student/assignments.html', {
        **_base_ctx(request), 'rows': rows, 'sf': sf, 'now': timezone.now(),
        'filter_tabs': [('all','Barchasi'),('pending','Kutilmoqda'),('submitted','Topshirildi'),('expired',"Muddati o'tdi")],})


@role_required('talaba')
def student_submit(request, pk):
    a = get_object_or_404(Assignment, pk=pk, status='active')
    s = request.user
    if a.is_expired:
        messages.error(request, "Topshiriq muddati o'tdi."); return redirect('core:student_assignments')
    if not (s.group and a.groups.filter(pk=s.group.pk).exists()):
        messages.error(request, "Bu topshiriq sizga tegishli emas."); return redirect('core:student_assignments')
    # Tanlangan talabalar - agar belgilangan bo'lsa va siz ro'yxatda yo'q bo'lsangiz
    if a.allowed_students.exists() and not a.allowed_students.filter(pk=s.pk).exists():
        messages.error(request, "Bu topshiriq sizga yuborilmagan."); return redirect('core:student_assignments')
    if Submission.objects.filter(assignment=a, student=s).exists():
        messages.warning(request, "Allaqachon topshirgansiz."); return redirect('core:student_assignments')
    # Race condition himoyasi: unique_together constraint IntegrityError ni ushlaymiz

    # Bir vaqtda faqat bir joydan kirish: session key tekshiruvi
    session_key = f'test_active_{a.pk}_{s.pk}'
    if request.method == 'GET':
        request.session[session_key] = request.session.session_key

    if a.assignment_type == 'test':
        return _handle_test(request, a, s)
    else:
        return _handle_written_or_file(request, a, s)


def _handle_test(request, a, s):
    import random
    from assessment.models import TopicQuota

    # Savol tartibini session da saqlaymiz — GET va POST da bir xil bo'lishi uchun
    session_key = f'test_questions_{a.pk}_{s.pk}'

    # Savol tartibi: session da bor bo'lsa qayta ishlatiladi (sahifa yangilansa ham)
    q_ids = request.session.get(session_key, [])
    if q_ids:
        # Session da saqlangan tartibni tiklaymiz
        q_map = {q.pk: q for q in a.questions.filter(pk__in=q_ids)}
        questions = [q_map[pk] for pk in q_ids if pk in q_map]
        # Savollar topilmasa (session eskirgan) — qayta quramiz
        if not questions:
            q_ids = []

    if not q_ids:
        # Har bir talabaga unikal seed — deterministik aralash (hamma boshqacha ko'radi)
        seed = hash(f"{s.pk}_{a.pk}") % (2**31)
        rng  = random.Random(seed)

        quotas = {q.topic: q.count for q in a.topic_quotas.all()}
        if quotas:
            selected = []
            for topic, count in quotas.items():
                topic_qs = list(a.questions.filter(topic=topic))
                if a.shuffle_questions:
                    rng.shuffle(topic_qs)
                selected.extend(topic_qs[:count])
            no_topic_qs = list(a.questions.filter(topic=''))
            if a.shuffle_questions:
                rng.shuffle(no_topic_qs)
            selected.extend(no_topic_qs)
            questions = selected
            if a.shuffle_questions:
                rng.shuffle(questions)
        else:
            questions = list(a.questions.order_by('order'))
            if a.shuffle_questions:
                rng.shuffle(questions)
            if a.questions_per_student > 0:
                questions = questions[:a.questions_per_student]

    if request.method == 'POST':
        answers = {str(q.pk): request.POST.get(f'ans_{q.pk}','') for q in questions}
        tab_switches  = int(request.POST.get('tab_switches', 0))
        copy_count    = int(request.POST.get('copy_count', 0))
        paste_count   = int(request.POST.get('paste_count', 0))
        # Shubhali xatti-harakatlar tahlili
        sus_events = []
        if tab_switches > 5:
            sus_events.append({'type':'tab_switch','count':tab_switches,'severity':'high'})
        elif tab_switches > 2:
            sus_events.append({'type':'tab_switch','count':tab_switches,'severity':'medium'})
        if copy_count > 3:
            sus_events.append({'type':'copy','count':copy_count,'severity':'high'})
        if paste_count > 1:
            sus_events.append({'type':'paste','count':paste_count,'severity':'high'})
        time_taken    = int(request.POST.get('time_taken', 0))
        from assessment.services.ai_service import grade_test
        qwa = [{'text': q.text, 'correct_answer': q.correct_answer,
                'student_answer': answers.get(str(q.pk),'')} for q in questions]
        r   = grade_test(a.title, qwa)
        try:
            sub = Submission.objects.create(
                assignment=a, student=s, test_answers=answers,
                tab_switches=tab_switches, time_taken_seconds=time_taken,
                ai_score=r['score'], ai_feedback=r['feedback'],
                ai_graded_at=timezone.now(), status='submitted',
                ip_address=request.META.get('REMOTE_ADDR'),
                suspicious_events=sus_events)
        except Exception:
            # Race condition — parallel so'rov allaqachon yaratgan
            messages.warning(request, "Allaqachon topshirilgan.")
            return redirect('core:student_assignments')
        messages.success(request, f"Test topshirildi! Dastlabki ball: {r['score']:.0f}/100")
        return redirect('core:student_result', pk=sub.pk)
    # Savol tartibini session da saqlaymiz
    request.session[session_key] = [q.pk for q in questions]
    request.session.modified = True

    q_json = json.dumps([{
        'id': q.pk, 'text': q.text,
        'image': q.image.url if q.image else None,
        'options': [{'k':'A','t':q.option_a},{'k':'B','t':q.option_b},{'k':'C','t':q.option_c}]
                    + ([{'k':'D','t':q.option_d}] if q.option_d else []),
        'accessible': q.is_accessible,
    } for q in questions])
    return render(request, 'student/take_test.html', {
        **_base_ctx(request), 'assignment': a, 'questions': questions,
        'q_json': q_json, 'duration': a.duration_minutes * 60,
        'accessible': s.is_accessible,
    })


def _handle_written_or_file(request, a, s):
    if request.method == 'POST':
        # Fayl validatsiyasi
        if 'file' in request.FILES:
            uploaded = request.FILES['file']
            # Server-side MIME tekshiruvi (brauzer headeriga ishonmaymiz)
            import os as _os
            # Fayl nomi xavfsizligi (path traversal himoyasi)
            safe_name    = _os.path.basename(uploaded.name)
            uploaded.name = safe_name

            # Fayl baytlaridan haqiqiy MIME ni aniqlaymiz
            try:
                import filetype as _ft
                header_bytes = uploaded.read(261)
                uploaded.seek(0)
                kind = _ft.guess(header_bytes)
                if kind:
                    actual_mime = kind.mime
                else:
                    actual_mime = uploaded.content_type or 'application/octet-stream'
            except ImportError:
                actual_mime = uploaded.content_type or ''

            ALLOWED_MIMES = {
                'application/pdf',
                'application/msword',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'application/zip',
                'application/x-zip-compressed',
                'image/jpeg', 'image/png', 'image/gif', 'image/webp',
                'text/plain',
            }
            is_allowed = (
                actual_mime in ALLOWED_MIMES
                or actual_mime.startswith('image/')
                or actual_mime.startswith('application/vnd.')
            )
            if not is_allowed:
                from django.contrib import messages as _msg
                _msg.error(request, f"Ruxsat etilmagan fayl turi: {actual_mime}. "
                                    f"Qabul qilinadi: PDF, Word, Excel, ZIP, rasm.")
                return render(request, 'student/submit_assignment.html', {
                    **_base_ctx(request), 'assignment': a, 'accessible': s.is_accessible})
            # Hajm tekshiruvi
            max_bytes = a.max_file_size_mb * 1024 * 1024
            if uploaded.size > max_bytes:
                from django.contrib import messages as msg
                msg.error(request, f"Fayl hajmi {a.max_file_size_mb}MB dan oshmasligi kerak. "
                          f"Sizning faylingiz: {uploaded.size//1024//1024:.1f}MB")
                return render(request, 'student/submit_assignment.html', {
                    **_base_ctx(request), 'assignment': a, 'accessible': s.is_accessible})
            # Tur tekshiruvi
            if a.allowed_file_types:
                import os
                ext = os.path.splitext(uploaded.name)[1].lower()
                if ext not in [t.lower() for t in a.allowed_file_types]:
                    from django.contrib import messages as msg
                    allowed = ', '.join(a.allowed_file_types)
                    msg.error(request, f"Ruxsat etilgan fayl turlari: {allowed}. Siz yuklagan: {ext}")
                    return render(request, 'student/submit_assignment.html', {
                        **_base_ctx(request), 'assignment': a, 'accessible': s.is_accessible})

        # Race condition himoyasi — duplicate submission tekshiruvi
        if Submission.objects.filter(assignment=a, student=s).exists():
            messages.warning(request, "Allaqachon topshirilgan.")
            return redirect('core:student_assignments')

        sub = Submission(assignment=a, student=s,
            text_answer=request.POST.get('text_answer',''),
            ip_address=request.META.get('REMOTE_ADDR'),
            status='submitted')
        if 'file' in request.FILES:
            sub.uploaded_file = request.FILES['file']
        # AI baholash background thread da — talaba darhol javob oladi
        if a.assignment_type == 'written' and sub.text_answer:
            sub.ai_feedback = "AI baholanyapti..."  # placeholder
        # Submission ni avval saqlaymiz (race condition oldini olish)
        try:
            sub.save()
        except Exception:
            messages.warning(request, "Allaqachon topshirilgan.")
            return redirect('core:student_assignments')

        # AI baholashni orqada ishga tushiramiz
        if a.assignment_type == 'written' and sub.text_answer:
            import threading
            def _ai_grade_async(sub_pk, title, subject, desc, answer):
                try:
                    from assessment.services.ai_service import grade_written
                    from assessment.models import Submission as _Sub
                    from django.utils import timezone as _tz
                    r = grade_written(title, subject, desc, answer)
                    _Sub.objects.filter(pk=sub_pk).update(
                        ai_score      = r.get('score', 0),
                        ai_feedback   = r.get('feedback', ''),
                        ai_strengths  = r.get('strengths', []),
                        ai_improvements=r.get('improvements', []),
                        ai_graded_at  = _tz.now(),
                    )
                except Exception as e:
                    logger.error(f"AI grading async error: {e}")
            t = threading.Thread(target=_ai_grade_async,
                args=(sub.pk, a.title, a.subject.name, a.description, sub.text_answer),
                daemon=True)
            t.start()

        messages.success(request, "Topshiriq yuborildi!")
        return redirect('core:student_result', pk=sub.pk)
    return render(request, 'student/submit_assignment.html', {
        **_base_ctx(request), 'assignment': a, 'accessible': s.is_accessible})


@role_required('talaba')
def student_result(request, pk):
    sub = get_object_or_404(Submission, pk=pk, student=request.user)
    if request.method == 'POST' and request.POST.get('action') == 'feedback':
        msg = request.POST.get('message','').strip()
        if msg:
            Feedback.objects.create(submission=sub, student=request.user, message=msg)
            from assessment.services.ai_service import send_notification
            send_notification(sub.assignment.teacher,
                f"Yangi shikoyat: {sub.student.full_name}",
                f"'{sub.assignment.title}' uchun shikoyat: {msg[:80]}", 'warning')
            messages.success(request, "Shikoyat yuborildi.")
            return redirect('core:student_result', pk=pk)
    q_results = []
    if sub.assignment.assignment_type == 'test':
        for q in Question.objects.filter(assignment=sub.assignment).order_by('order'):
            sa = sub.test_answers.get(str(q.pk),'—')
            q_results.append({'q': q, 'sa': sa, 'ok': sa.upper()==q.correct_answer.upper()})
    return render(request, 'student/result.html', {
        **_base_ctx(request), 'sub': sub, 'q_results': q_results,
        'existing_fb': sub.feedbacks.order_by('-created_at').first(),
    })


@role_required('talaba')
def student_progress(request):
    s    = request.user
    subs = Submission.objects.filter(student=s).select_related('assignment__subject').order_by('-submitted_at')
    by_subject = {}
    for sub in subs:
        name = sub.assignment.subject.name
        if name not in by_subject: by_subject[name] = []
        sc = sub.final_score if sub.final_score is not None else sub.ai_score
        if sc is not None: by_subject[name].append(sc)
    stats = [{'name': n, 'count': len(sc), 'avg': round(sum(sc)/len(sc),1),
              'max': max(sc), 'min': min(sc)} for n, sc in by_subject.items() if sc]
    overall = subs.filter(final_score__isnull=False).aggregate(a=Avg('final_score'))['a']
    return render(request, 'student/progress.html', {
        **_base_ctx(request),
        'submissions': subs[:20], 'stats': stats,
        'overall_avg': round(overall,1) if overall else None,
        'total_graded': subs.filter(status='graded').count(),
        'total_submitted': subs.count(),
        'grade_bands': [
            ('5', 'bg-green-100 text-green-800',  '86', '100'),
            ('4', 'bg-blue-100 text-blue-800',    '71', '85'),
            ('3', 'bg-yellow-100 text-yellow-800','56', '70'),
            ('2', 'bg-orange-100 text-orange-800','41', '55'),
            ('1', 'bg-red-100 text-red-700',      '0',  '40'),
        ],
    })


# ─── Notifications ─────────────────────────────────────────────
@login_required

@login_required
def toggle_accessible(request, pk):
    """Talabani inklyuziv / oddiy ta'limga o'tkazish (kafedra mudiri va admin)"""
    if not (request.user.is_admin or request.user.is_kafedra):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden()
    
    from accounts.models import User as _User
    student = get_object_or_404(_User, pk=pk, role='talaba')
    
    # Kafedra mudiri faqat o'z kafedra talabasini o'zgartira oladi
    if request.user.is_kafedra and student.department != request.user.department:
        messages.error(request, "Ruxsat yo'q.")
        return redirect('core:kafedra_students')
    
    student.is_accessible = not student.is_accessible
    student.save(update_fields=['is_accessible'])
    
    status = "Inklyuziv ta'lim" if student.is_accessible else "Oddiy ta'lim"
    messages.success(request, f"{student.full_name} — {status} ga o'tkazildi.")
    
    if request.user.is_admin:
        return redirect('core:admin_users')
    return redirect('core:kafedra_students')


def notifications(request):
    notifs = Notification.objects.filter(recipient=request.user).order_by('-created_at')
    notifs.filter(is_read=False).update(is_read=True)
    return render(request, 'base/notifications.html', {**_base_ctx(request), 'notifications': notifs})


@login_required
def mark_read(request, pk):
    Notification.objects.filter(pk=pk, recipient=request.user).update(is_read=True)
    return JsonResponse({'ok': True})


# ══════════════════════════════════════════════════════════════
# SAVOL BANKI (Question Bank)
# ══════════════════════════════════════════════════════════════
@role_required('oqituvchi')
def teacher_question_banks(request):
    """Savollar banki ro'yxati"""
    t     = request.user
    banks = QuestionBank.objects.filter(teacher=t).select_related('subject').annotate(
        qcount=Count('bank_questions')
    ).order_by('-updated_at')

    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'create':
            from assessment.models import QuestionBank as QB
            QB.objects.create(
                teacher=t,
                subject_id=request.POST['subject'],
                title=request.POST['title'],
                description=request.POST.get('description', ''),
            )
            messages.success(request, "Savol banki yaratildi.")
        elif act == 'delete':
            from assessment.models import QuestionBank as QB
            QB.objects.filter(pk=request.POST['id'], teacher=t).delete()
            messages.success(request, "O'chirildi.")
        return redirect('core:teacher_question_banks')

    my_subjects = SubjectTeacher.objects.filter(teacher=t).select_related('subject')
    return render(request, 'teacher/question_banks.html', {
        **_base_ctx(request),
        'banks': banks,
        'my_subjects': my_subjects,
    })


@role_required('oqituvchi')
def teacher_bank_detail(request, pk):
    """Bank ichidagi savollar"""
    from assessment.models import QuestionBank as QB, BankQuestion
    bank = get_object_or_404(QB, pk=pk, teacher=request.user)

    if request.method == 'POST':
        act = request.POST.get('action')

        if act == 'add_question':
            q = BankQuestion.objects.create(
                bank=bank,
                text=request.POST['text'],
                topic=request.POST.get('topic', ''),
                correct_answer=request.POST['correct_answer'],
                option_a=request.POST['option_a'],
                option_b=request.POST['option_b'],
                option_c=request.POST['option_c'],
                option_d=request.POST.get('option_d', ''),
                difficulty=request.POST.get('difficulty', 'medium'),
            )
            if 'image' in request.FILES:
                q.image = request.FILES['image']
                q.save()
            messages.success(request, "Savol qo'shildi.")

        elif act == 'delete_question':
            BankQuestion.objects.filter(pk=request.POST['qid'], bank=bank).delete()
            messages.success(request, "O'chirildi.")

        elif act == 'import_docx':
            _import_bank_docx(request, bank)

        elif act == 'check_duplicates':
            _check_bank_duplicates(bank)
            messages.success(request, "AI takroriy savollarni tekshirdi.")

        elif act == 'use_in_assignment':
            # Tanlangan savollarni topshiriqqa ko'chirish
            assignment_id = request.POST.get('assignment_id')
            selected_ids  = request.POST.getlist('selected_questions')
            if assignment_id and selected_ids:
                assignment = get_object_or_404(Assignment, pk=assignment_id, teacher=request.user)
                for qid in selected_ids:
                    bq = get_object_or_404(BankQuestion, pk=qid, bank=bank)
                    Question.objects.create(
                        assignment=assignment,
                        text=bq.text,
                        topic=bq.topic,
                        correct_answer=bq.correct_answer,
                        option_a=bq.option_a,
                        option_b=bq.option_b,
                        option_c=bq.option_c,
                        option_d=bq.option_d,
                        order=assignment.questions.count() + 1,
                    )
                    bq.use_count += 1
                    bq.save()
                messages.success(request, f"{len(selected_ids)} ta savol topshiriqqa qo'shildi.")
                return redirect('core:teacher_edit_assignment', pk=assignment_id)

        return redirect('core:teacher_bank_detail', pk=pk)

    # Topshiriqlar ro'yxati (savollarni ularga ko'chirish uchun)
    my_assignments = Assignment.objects.filter(
        teacher=request.user, subject=bank.subject, status='draft'
    ).order_by('-created_at')

    topic_filter = request.GET.get('topic', '')
    diff_filter  = request.GET.get('diff', '')
    questions = bank.bank_questions.all()
    if topic_filter:
        questions = questions.filter(topic__icontains=topic_filter)
    if diff_filter:
        questions = questions.filter(difficulty=diff_filter)

    topics = bank.bank_questions.values_list('topic', flat=True).distinct().exclude(topic='')

    return render(request, 'teacher/bank_detail.html', {
        **_base_ctx(request),
        'bank': bank,
        'questions': questions,
        'my_assignments': my_assignments,
        'topics': topics,
        'topic_filter': topic_filter,
        'diff_filter': diff_filter,
    })


def _import_bank_docx(request, bank):
    """Word fayldan bank savollarini import qilish"""
    import tempfile, os
    from assessment.models import BankQuestion
    f = request.FILES.get('docx_file')
    if not f:
        return
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        for chunk in f.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name
    try:
        from docx import Document
        doc = Document(tmp_path)
        n = 0
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 3 and cells[1]:
                    qdata = _parse_question_from_cells(cells)
                    if qdata:
                        BankQuestion.objects.create(
                            bank=bank,
                            text=qdata['text'],
                            correct_answer=qdata['correct_answer'],
                            option_a=qdata['option_a'],
                            option_b=qdata['option_b'],
                            option_c=qdata['option_c'],
                            option_d=qdata['option_d'],
                        )
                        n += 1
        messages.success(request, f"{n} ta savol import qilindi.")
    except Exception as e:
        messages.error(request, f"Import xatosi: {e}")
    finally:
        os.unlink(tmp_path)


def _check_bank_duplicates(bank):
    """AI yordamida takroriy savollarni aniqlash"""
    from assessment.models import BankQuestion
    from assessment.services.ai_service import get_client
    import json

    questions = list(bank.bank_questions.all())
    if len(questions) < 2:
        return

    client = get_client()
    if not client:
        # Oddiy matn o'xshashligi
        seen = {}
        for q in questions:
            key = q.text[:50].lower().strip()
            if key in seen:
                q.is_duplicate = True
                q.duplicate_note = f"#{seen[key]} savol bilan o'xshash"
                q.save()
            else:
                seen[key] = q.pk
        return

    # AI bilan tekshirish (25 tadan oshsa bo'laklash)
    batch = questions[:25]
    q_list = "\n".join([f"{i+1}. {q.text[:100]}" for i, q in enumerate(batch)])
    prompt = f"""Quyidagi test savollarini ko'rib chiqing va takroriy yoki juda o'xshash savollarni aniqlang.

Savollar:
{q_list}

Faqat JSON qaytaring:
{{"duplicates": [
  {{"index": 3, "similar_to": 1, "note": "bir xil savol"}},
  {{"index": 7, "similar_to": 5, "note": "juda o'xshash"}}
]}}
Takroriy savollar yo'q bo'lsa: {{"duplicates": []}}"""

    try:
        r = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )
        text = r.content[0].text.strip()
        # Mustahkam JSON parsing
        data = None
        try:
            import re as _re
            m = _re.search(r'\{.*\}', text, _re.DOTALL)
            if m:
                data = json.loads(m.group())
        except Exception:
            pass
        if data:
            for dup in data.get('duplicates', []):
                idx = dup.get('index', 0) - 1
                if 0 <= idx < len(batch):
                    batch[idx].is_duplicate = True
                    batch[idx].duplicate_note = dup.get('note', 'Takroriy')
                    batch[idx].save()
    except Exception as e:
        logger.warning(f"Bank duplicate tekshiruv xato: {e}")


# ══════════════════════════════════════════════════════════════
# MAVZU KVOTA (Topic Quota)
# ══════════════════════════════════════════════════════════════
@role_required('oqituvchi')
def teacher_topic_quota(request, pk):
    """Topshiriqdagi mavzu bo'yicha savol kvotasini belgilash"""
    from assessment.models import TopicQuota
    assignment = get_object_or_404(Assignment, pk=pk, teacher=request.user)

    if request.method == 'POST':
        act = request.POST.get('action')
        if act == 'save_quotas':
            TopicQuota.objects.filter(assignment=assignment).delete()
            topics  = request.POST.getlist('topic')
            counts  = request.POST.getlist('count')
            for topic, count in zip(topics, counts):
                topic = topic.strip()
                if topic and count:
                    try:
                        TopicQuota.objects.create(
                            assignment=assignment,
                            topic=topic,
                            count=int(count),
                        )
                    except ValueError:
                        pass
            messages.success(request, "Kvotalar saqlandi.")
        return redirect('core:teacher_topic_quota', pk=pk)

    # Mavjud mavzular (savollardan)
    existing_topics = list(
        assignment.questions.values_list('topic', flat=True)
        .distinct().exclude(topic='')
    )
    existing_quotas = {q.topic: q.count for q in assignment.topic_quotas.all()}

    # Sillabus mavzulari
    syllabus = assignment.subject.syllabi.order_by('-uploaded_at').first()
    syllabus_topics = syllabus.topics if syllabus else []

    all_topics = list(set(existing_topics + syllabus_topics))

    return render(request, 'teacher/topic_quota.html', {
        **_base_ctx(request),
        'assignment': assignment,
        'all_topics': all_topics,
        'existing_quotas': existing_quotas,
        'total_questions': assignment.questions.count(),
        'syllabus_topics': syllabus_topics,
    })


# ══════════════════════════════════════════════════════════════
# GURUH EXCEL IMPORT
# ══════════════════════════════════════════════════════════════
@role_required('kafedra_mudiri', 'admin')
def import_students_excel(request):
    """
    Universal Excel import — barcha turdagi foydalanuvchilar uchun.
    Ustunlar: Familya | Ism | Login* | Parol | ID | Guruh/Kafedra | ROL*
    Rol qiymatlari: talaba, oqituvchi, kafedra_mudiri, admin
    """
    if request.method == 'POST' and 'excel_file' in request.FILES:
        try:
            import openpyxl
            from accounts.models import User

            wb  = openpyxl.load_workbook(request.FILES['excel_file'])
            ws  = wb.active
            dept = request.user.department

            # Ruxsat etilgan rollar (kim kim qo'sha oladi)
            if request.user.is_admin:
                allowed_roles = {'talaba', 'oqituvchi', 'kafedra_mudiri', 'admin'}
            elif request.user.is_kafedra:
                allowed_roles = {'talaba', 'oqituvchi'}
            else:
                messages.error(request, "Ruxsat yo'q.")
                return redirect('core:kafedra_students')

            # Rol nomlari mapping (foydalanuvchi yozishi mumkin bo'lgan variantlar)
            role_map = {
                'talaba':          'talaba',
                'student':         'talaba',
                "o'qituvchi":     "oqituvchi",
                'oqituvchi':       'oqituvchi',
                'teacher':         'oqituvchi',
                'kafedra_mudiri':  'kafedra_mudiri',
                'kafedra mudiri':  'kafedra_mudiri',
                'kafedra':         'kafedra_mudiri',
                'dean':            'kafedra_mudiri',
                'admin':           'admin',
                'administrator':   'admin',
            }

            created = 0; skipped = 0; errors = []
            role_counts = {}

            for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if not any(row):
                    continue
                try:
                    # Ustunlar yetishmasa bo'sh string sifatida ishlatamiz (IndexError oldini olish)
                    def _cell(idx, default=''):
                        try:
                            v = row[idx]
                            return str(v).strip() if v is not None else default
                        except IndexError:
                            return default

                    last_name   = _cell(0)
                    first_name  = _cell(1)
                    username    = _cell(2)
                    password    = _cell(3, settings.DEFAULT_STUDENT_PASSWORD) or settings.DEFAULT_STUDENT_PASSWORD
                    extra_id    = _cell(4)
                    group_dept  = _cell(5)
                    role_raw    = _cell(6, 'talaba').lower() or 'talaba'

                    if not username:
                        errors.append(f"{i}-qator: Login (username) bo'sh")
                        continue

                    # Rol aniqlash
                    role = role_map.get(role_raw, 'talaba')

                    if role not in allowed_roles:
                        errors.append(f"{i}-qator: '{role_raw}' rolini qo'shish huquqingiz yo'q")
                        continue

                    # Mavjud foydalanuvchini guruhga biriktirib yuboramiz (kafedra mudiri ehtiyojini)
                    existing = User.objects.filter(username=username).first()
                    if existing:
                        # Faqat shu kafedraga tegishli bo'lsa — guruhini yangilaymiz
                        # Boshqa kafedraniki bo'lsa, o'tkazib yuboramiz (xavfsizlik)
                        if dept and existing.department and existing.department != dept:
                            errors.append(f"{i}-qator: '{username}' boshqa kafedrada ({existing.department.name})")
                            skipped += 1
                            continue
                        # Guruhsiz yoki shu kafedra ichida — yangilaymiz
                        if not existing.department:
                            existing.department = dept
                        # Guruhni topib biriktiramiz
                        if existing.role == 'talaba' and group_dept:
                            grp = Group.objects.filter(name__iexact=group_dept, department=dept).first()
                            if grp:
                                existing.group = grp
                        if extra_id and existing.role == 'talaba':
                            existing.student_id = extra_id
                        existing.save()
                        skipped += 1
                        continue

                    # Guruh yoki kafedra aniqlash
                    group      = None
                    user_dept  = dept

                    if role in ('talaba', 'oqituvchi'):
                        if group_dept:
                            group = Group.objects.filter(
                                name__iexact=group_dept,
                                department=dept
                            ).first()
                            if not group and dept is None:
                                # Admin uchun istalgan departmentdan qidirish
                                group = Group.objects.filter(name__iexact=group_dept).first()
                        # Kafedra aniqlashtirish
                        if group:
                            user_dept = group.department
                        elif group_dept:
                            user_dept = Department.objects.filter(
                                name__icontains=group_dept
                            ).first() or dept
                    elif role == 'kafedra_mudiri':
                        if group_dept:
                            user_dept = Department.objects.filter(
                                name__icontains=group_dept
                            ).first() or dept

                    # Foydalanuvchi yaratish
                    create_kwargs = dict(
                        username=username,
                        password=password,
                        first_name=first_name,
                        last_name=last_name,
                        role=role,
                        department=user_dept,
                    )
                    if role == 'talaba':
                        create_kwargs['student_id'] = extra_id
                        create_kwargs['group'] = group
                    elif role == 'oqituvchi':
                        create_kwargs['group'] = None

                    u = User.objects.create_user(**create_kwargs)
                    created += 1
                    role_counts[role] = role_counts.get(role, 0) + 1

                except Exception as e:
                    errors.append(f"{i}-qator: {str(e)[:80]}")

            # Natija xabari
            parts = []
            role_labels = {
                'talaba': 'talaba', 'oqituvchi': "o'qituvchi",
                'kafedra_mudiri': 'kafedra mudiri', 'admin': 'admin'
            }
            for r, cnt in role_counts.items():
                parts.append(f"{cnt} {role_labels.get(r, r)}")

            if created:
                messages.success(request, f"✅ Import muvaffaqiyatli: {', '.join(parts)}")
            if skipped:
                messages.info(request, f"{skipped} ta foydalanuvchi allaqachon mavjud — o'tkazib yuborildi.")
            for err in errors[:5]:
                messages.warning(request, err)
            if not created and not skipped:
                messages.error(request, "Hech qanday foydalanuvchi import qilinmadi.")

        except Exception as e:
            messages.error(request, f"Excel import xatosi: {e}")

    if request.user.is_admin:
        return redirect('core:admin_users')
    return redirect('core:kafedra_students')


@login_required
def download_student_template(request):
    """Talabalarni import qilish uchun Excel shablon"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from io import BytesIO

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Talabalar"

    headers = ["Familya", "Ism", "Login (username)*", "Parol", "Talaba ID", "Guruh nomi"]
    header_fill = PatternFill("solid", fgColor="00236F")
    header_font = Font(bold=True, color="FFFFFF")
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # Sample data
    samples = [
        ["Rahimov", "Abdulloh", "abdulloh_r", settings.DEFAULT_STUDENT_PASSWORD, "STU001", "IF-201"],
        ["Tursunova", "Barno",   "barno_t",    settings.DEFAULT_STUDENT_PASSWORD, "STU002", "IF-201"],
    ]
    for r, row in enumerate(samples, 2):
        for c, val in enumerate(row, 1):
            ws.cell(row=r, column=c, value=val)

    col_widths = [15, 12, 20, 15, 12, 12]
    from openpyxl.utils import get_column_letter
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws2 = wb.create_sheet("Ko'rsatmalar")
    ws2['A1'] = "Ko'rsatmalar:"
    ws2['A1'].font = Font(bold=True)
    notes = [
        "1. Login (username) — MAJBURIY, unikal bo'lishi kerak",
        "2. Parol — bo'sh qolsa settings.DEFAULT_STUDENT_PASSWORD ishlatiladi",
        "3. Guruh nomi — tizimda mavjud guruh nomini kiriting (masalan: IF-201)",
        "4. 1-qator sarlavha, 2-qatordan boshlab ma'lumot kiriting",
        "5. Allaqachon mavjud username lar o'tkazib yuboriladi",
    ]
    for i, note in enumerate(notes, 2):
        ws2.cell(row=i, column=1, value=note)
    ws2.column_dimensions['A'].width = 60

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    resp = HttpResponse(buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp['Content-Disposition'] = 'attachment; filename="talabalar_import_shablon.xlsx"'
    return resp



@login_required
def tts_audio(request):
    """O'zbek tilida TTS — edge-tts (Microsoft, bepul, uz-UZ-SardorNeural ovozi)"""
    import hashlib, os
    from django.http import FileResponse, HttpResponse

    text = request.GET.get('text', '').strip()
    if not text:
        return HttpResponse(status=400)

    # Kesh — bir xil matn qayta yaratilmaydi
    text_hash  = hashlib.md5(text.encode('utf-8')).hexdigest()
    cache_dir  = os.path.join(settings.MEDIA_ROOT, 'tts')
    os.makedirs(cache_dir, exist_ok=True)
    cache_file = os.path.join(cache_dir, f'{text_hash}.mp3')

    if not os.path.exists(cache_file):
        try:
            import asyncio, edge_tts

            async def _generate():
                # uz-UZ-SardorNeural = erkak, uz-UZ-MadinaNeural = ayol
                comm = edge_tts.Communicate(
                    text=text,
                    voice='uz-UZ-SardorNeural',
                    rate='-5%',
                )
                await comm.save(cache_file)

            asyncio.run(_generate())

        except ImportError:
            logger.warning("edge-tts kutubxonasi o'rnatilmagan")
            return HttpResponse(status=404)
        except Exception as e:
            logger.error(f"TTS xato: {e}")
            if os.path.exists(cache_file):
                os.remove(cache_file)
            return HttpResponse(status=500)

    try:
        resp = FileResponse(open(cache_file, 'rb'), content_type='audio/mpeg')
        resp['Cache-Control'] = 'public, max-age=86400'
        return resp
    except Exception:
        return HttpResponse(status=404)


@login_required
def admin_edit_user(request, pk):
    """Admin foydalanuvchini tahrirlash"""
    if not request.user.is_admin:
        return redirect('core:dashboard')
    from accounts.models import User as _User
    from core.models import Department, Group
    u = get_object_or_404(_User, pk=pk)
    if request.method == 'POST':
        u.first_name    = request.POST.get('first_name', u.first_name)
        u.last_name     = request.POST.get('last_name', u.last_name)
        u.email         = request.POST.get('email', u.email)
        u.is_accessible = 'is_accessible' in request.POST
        dept_pk = request.POST.get('department')
        if dept_pk:
            u.department = Department.objects.filter(pk=dept_pk).first()
        if request.POST.get('password'):
            u.set_password(request.POST['password'])
        u.save()
        messages.success(request, f"{u.full_name} yangilandi.")
        return redirect('core:admin_users')
    return render(request, 'admin_panel/users.html', {
        **_base_ctx(request),
        'edit_user': u,
    })


@login_required
def kafedra_edit_student(request, pk):
    """Kafedra mudiri talabani tahrirlash"""
    if not (request.user.is_kafedra or request.user.is_admin):
        return redirect('core:dashboard')
    from accounts.models import User as _User
    from core.models import Group
    student = get_object_or_404(_User, pk=pk, role='talaba')
    if request.user.is_kafedra and student.department != request.user.department:
        messages.error(request, "Ruxsat yo'q.")
        return redirect('core:kafedra_students')
    if request.method == 'POST':
        student.first_name    = request.POST.get('first_name', student.first_name)
        student.last_name     = request.POST.get('last_name', student.last_name)
        student.student_id    = request.POST.get('student_id', student.student_id)
        student.is_accessible = 'is_accessible' in request.POST
        grp_pk = request.POST.get('group')
        if grp_pk:
            student.group = Group.objects.filter(pk=grp_pk).first()
        if request.POST.get('password'):
            student.set_password(request.POST['password'])
        student.save()
        messages.success(request, f"{student.full_name} yangilandi.")
        return redirect('core:kafedra_students')
    return redirect('core:kafedra_students')


def error_404(request, exception=None):
    return render(request, '404.html', status=404)


def error_500(request):
    return render(request, '500.html', status=500)

@login_required
def submit_appeal(request, sub_pk):
    """Talaba bahoga e'tiroz bildiradi"""
    if not request.user.is_student:
        return redirect('core:student_dashboard')
    sub = get_object_or_404(Submission, pk=sub_pk, student=request.user)

    if request.method == 'POST':
        reason = request.POST.get('reason', '').strip()
        if len(reason) < 20:
            messages.error(request, "E'tiroz sababini kamida 20 belgi yozing.")
        elif sub.appeal_status not in ('none', 'rejected'):
            messages.warning(request, "E'tiroz allaqachon yuborilgan.")
        else:
            sub.appeal_status = 'pending'
            sub.appeal_reason = reason
            sub.appealed_at   = timezone.now()
            sub.save(update_fields=['appeal_status', 'appeal_reason', 'appealed_at'])
            Notification.objects.create(
                recipient=sub.assignment.teacher,
                title="Bahoga e'tiroz",
                message=f"{sub.student.full_name} '{sub.assignment.title}' bahosiga e'tiroz bildirdi.",
                notification_type='feedback',
                link=f'/dashboard/teacher/grade-book/?assignment={sub.assignment_id}'
            )
            messages.success(request, "E'tirozingiz yuborildi.")
        return redirect('core:student_progress')

    return render(request, 'student/submit_appeal.html', {**_base_ctx(request), 'sub': sub})


@login_required
def review_appeal(request, sub_pk):
    """O'qituvchi e'tirozni ko'rib chiqadi"""
    if not request.user.is_teacher:
        return redirect('core:teacher_dashboard')
    sub = get_object_or_404(Submission, pk=sub_pk, assignment__teacher=request.user)

    if request.method == 'POST':
        decision  = request.POST.get('decision')
        response  = request.POST.get('response', '').strip()
        new_score = request.POST.get('new_score', '').strip()

        if decision == 'accept':
            sub.appeal_status   = 'accepted'
            sub.appeal_response = response
            if new_score:
                try:
                    sub.final_score = float(new_score)
                    sub.graded_by   = request.user
                    sub.graded_at   = timezone.now()
                except ValueError:
                    pass
            sub.save()
            Notification.objects.create(
                recipient=sub.student,
                title="E'tiroz qabul qilindi",
                message=f"'{sub.assignment.title}' bahosi: {sub.final_score}",
                notification_type='graded',
            )
            messages.success(request, "E'tiroz qabul qilindi.")
        elif decision == 'reject':
            sub.appeal_status   = 'rejected'
            sub.appeal_response = response
            sub.save()
            Notification.objects.create(
                recipient=sub.student,
                title="E'tiroz rad etildi",
                message=response[:200],
                notification_type='graded',
            )
            messages.info(request, "E'tiroz rad etildi.")
        return redirect('core:teacher_grade_book')

    return render(request, 'teacher/review_appeal.html', {**_base_ctx(request), 'sub': sub})